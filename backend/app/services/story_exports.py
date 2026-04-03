from __future__ import annotations

import re
from io import BytesIO
from pathlib import PurePosixPath

from docx import Document
from sqlalchemy.orm import Session

from app.db import AssetKind, SessionAsset, StorySession
from app.repositories import SessionAssetRepository
from app.services.assets import AssetSessionNotFoundError, SessionAssetService
from app.storage import ObjectStorageService, StorageObjectLocation
from app.storage.service import ObjectNotFoundError, StorageError

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_DOCX_EXPORT_ID = "final-manuscript"
_HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+)$")


class StoryExportError(Exception):
    """Base error for story export failures."""


class StoryExportUnavailableError(StoryExportError):
    """Raised when the accepted story text is not ready for export."""


class StoryDocxExportService:
    def __init__(
        self,
        session: Session,
        *,
        object_storage: ObjectStorageService,
    ) -> None:
        self._session = session
        self._assets = SessionAssetRepository(session)
        self._asset_service = SessionAssetService(session)
        self._object_storage = object_storage

    def ensure_docx_asset(self, session_id: str) -> SessionAsset:
        story_session = self._session.get(StorySession, session_id)
        if story_session is None:
            raise AssetSessionNotFoundError(f"session {session_id!r} was not found")

        story_text_asset = self._assets.get_latest_ready(
            session_id,
            asset_kinds=(AssetKind.STORY_TEXT,),
        )
        if story_text_asset is None:
            raise StoryExportUnavailableError(
                "The accepted manuscript is not ready for Word export yet."
            )

        current_docx_asset = self._assets.get_latest_ready(
            session_id,
            asset_kinds=(AssetKind.STORY_DOCX,),
        )
        if _docx_matches_story_asset(current_docx_asset, story_text_asset):
            return current_docx_asset

        story_text = self._load_story_text(story_text_asset)
        docx_bytes = render_story_docx_bytes(
            title=_resolve_story_title(story_session),
            story_text=story_text,
            genre_label=getattr(getattr(story_session, "selected_genre", None), "label", None),
            tone_label=(
                getattr(getattr(story_session, "selected_tone_profile", None), "label", None)
            ),
        )
        location = self._object_storage.paths.export_asset(
            session_id=session_id,
            export_kind="docx",
            export_id=_DOCX_EXPORT_ID,
            extension="docx",
        )
        upload_metadata = self._object_storage.upload_bytes(
            location,
            docx_bytes,
            content_type=DOCX_MIME_TYPE,
        )
        asset_view = self._asset_service.upsert_asset_record(
            session_id=session_id,
            asset_kind=AssetKind.STORY_DOCX,
            storage_bucket=location.bucket,
            object_path=location.key,
            mime_type=DOCX_MIME_TYPE,
            status=story_text_asset.status,
            composition_job_id=story_text_asset.composition_job_id,
            segment_index=story_text_asset.segment_index,
            byte_size=upload_metadata.size_bytes,
            metadata_json={
                "export_version": "story_docx_export.v1",
                "download_filename": _build_story_docx_filename(story_session),
                "source_story_asset_id": story_text_asset.id,
                "source_story_ready_at": (
                    story_text_asset.ready_at.isoformat()
                    if story_text_asset.ready_at is not None
                    else None
                ),
                "source_story_object_path": story_text_asset.object_path,
                "title": _resolve_story_title(story_session),
                "word_count": _count_words(story_text),
            },
        )
        asset = self._session.get(SessionAsset, asset_view.id)
        if asset is None:
            raise StoryExportError("The Word export metadata could not be reloaded.")
        return asset

    def _load_story_text(self, asset: SessionAsset) -> str:
        location = StorageObjectLocation(
            bucket=asset.storage_bucket,
            key=asset.object_path,
        )
        try:
            return self._object_storage.download_text(location)
        except (ObjectNotFoundError, StorageError) as exc:
            raise StoryExportUnavailableError(
                "The accepted manuscript could not be loaded for Word export."
            ) from exc


def render_story_docx_bytes(
    *,
    title: str,
    story_text: str,
    genre_label: str | None = None,
    tone_label: str | None = None,
) -> bytes:
    document = Document()
    document.core_properties.title = title
    document.add_heading(title, level=0)

    metadata_parts = [part for part in (genre_label, tone_label) if part]
    if metadata_parts:
        document.add_paragraph(" | ".join(metadata_parts))

    for block_kind, block_value in _iter_story_blocks(story_text):
        if block_kind == "heading":
            heading_level, heading_text = block_value
            document.add_heading(heading_text, level=min(heading_level, 4))
            continue
        document.add_paragraph(block_value)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _iter_story_blocks(story_text: str) -> list[tuple[str, tuple[int, str] | str]]:
    blocks: list[tuple[str, tuple[int, str] | str]] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        paragraph = " ".join(line.strip() for line in paragraph_lines if line.strip()).strip()
        paragraph_lines.clear()
        if paragraph:
            blocks.append(("paragraph", paragraph))

    for raw_line in story_text.splitlines():
        heading_match = _HEADING_PATTERN.match(raw_line.strip())
        if heading_match is not None:
            flush_paragraph()
            blocks.append(
                (
                    "heading",
                    (len(heading_match.group(1)), heading_match.group(2).strip()),
                )
            )
            continue

        if not raw_line.strip():
            flush_paragraph()
            continue

        paragraph_lines.append(raw_line)

    flush_paragraph()
    return blocks


def _docx_matches_story_asset(
    docx_asset: SessionAsset | None,
    story_text_asset: SessionAsset,
) -> bool:
    if docx_asset is None:
        return False

    metadata = docx_asset.metadata_json
    if not isinstance(metadata, dict):
        return False

    return metadata.get("source_story_asset_id") == story_text_asset.id


def _resolve_story_title(story_session: StorySession) -> str:
    working_title = (story_session.working_title or "").strip()
    return working_title or "Bedtime Story"


def _build_story_docx_filename(story_session: StorySession) -> str:
    title = _resolve_story_title(story_session)
    slug = re.sub(r"[^A-Za-z0-9._-]+", "-", title).strip(".-_")
    normalized_title = slug or "bedtime-story"
    return f"{normalized_title}.docx"


def _count_words(text: str) -> int:
    return len([part for part in text.split() if part.strip()])


def resolve_asset_filename(asset: SessionAsset) -> str:
    metadata = asset.metadata_json if isinstance(asset.metadata_json, dict) else {}
    metadata_filename = metadata.get("download_filename")
    if isinstance(metadata_filename, str) and metadata_filename.strip():
        return metadata_filename.strip()
    fallback_name = PurePosixPath(asset.object_path).name.strip()
    if fallback_name:
        return fallback_name
    extension = {
        AssetKind.STORY_DOCX: "docx",
        AssetKind.STORY_TEXT: "md",
        AssetKind.AUDIO_SEGMENT: "wav",
        AssetKind.FINAL_AUDIO: "mp3",
    }.get(asset.asset_kind, "bin")
    return f"{asset.asset_kind.value}.{extension}"
