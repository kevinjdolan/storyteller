from __future__ import annotations

from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Iterator

import pytest
from app.db import (
    AssetKind,
    AssetStatus,
    AudioJob,
    Base,
    CompositionJob,
    CompositionJobKind,
    CompositionSegment,
    CompositionSegmentAcceptanceState,
    JobStatus,
    NarrationPauseHint,
    NarrationSegment,
    NarrationSourceBoundaryKind,
    SessionAsset,
    StorySession,
    StorySetup,
)
from app.db.session import get_engine, get_session_factory
from app.main import create_app
from app.models import (
    BeatSheetGenerationInvocation,
    BeatSheetGenerationInvocationResult,
    BeatSheetGenerationStructuredOutput,
    BriefNormalizationInvocation,
    BriefNormalizationInvocationResult,
    BriefNormalizationStructuredOutput,
    CharacterGenerationInvocation,
    CharacterGenerationInvocationResult,
    CharacterGenerationStructuredOutput,
    GeneratedBeatSheetBeat,
    GeneratedBeatSheetCandidate,
    GeneratedCharacterProfile,
    GeneratedCharacterSheetCandidate,
    GeneratedPitchCandidate,
    ModelCallOutcome,
    ModelUsageBucket,
    NormalizedBriefPreferences,
    PitchGenerationInvocation,
    PitchGenerationInvocationResult,
    PitchGenerationStructuredOutput,
    WorkflowStage,
    WorkflowStageState,
)
from app.services.catalog import CATALOG_FILE_PATH, load_catalog_document, seed_catalog
from app.services.model_usage import ModelUsageContext, SessionModelUsageService
from app.services.session_realtime import CompositionChunkCursor, SessionRealtimeService
from app.services.sessions import SessionService
from app.settings import get_settings
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import select
from tests.support.in_memory_storage import InMemoryObjectStorage


class StubBriefNormalizationAdapter:
    def __init__(self) -> None:
        self.model_id = "gemini-3.1-flash-lite"

    def normalize(
        self,
        invocation: BriefNormalizationInvocation,
    ) -> BriefNormalizationInvocationResult:
        return BriefNormalizationInvocationResult(
            invocation=invocation,
            structured_output=BriefNormalizationStructuredOutput(
                normalized_summary=(
                    "A harbor bedtime quest where each lantern return settles the night."
                ),
                normalized_preferences=NormalizedBriefPreferences(
                    protagonist_type="A child and an otter guide",
                    setting="a moonlit harbor",
                    emotional_goal="a calm return home",
                    constraint_notes=["End with the harbor settled and safe."],
                    bedtime_safety_concerns=["Keep every surprise quickly reassuring."],
                    candidate_motifs=["floating lanterns", "moonlit water"],
                ),
            ),
            raw_response={
                "stub": True,
                "usageMetadata": {
                    "promptTokenCount": 240,
                    "candidatesTokenCount": 72,
                    "totalTokenCount": 312,
                },
            },
        )

    def close(self) -> None:
        return None


class StubPitchGenerationAdapter:
    def __init__(self) -> None:
        self.model_id = "gemini-3.1-pro"

    def generate(
        self,
        invocation: PitchGenerationInvocation,
    ) -> PitchGenerationInvocationResult:
        return PitchGenerationInvocationResult(
            invocation=invocation,
            structured_output=PitchGenerationStructuredOutput(
                pitches=[
                    GeneratedPitchCandidate(
                        title="The Juniper Lake Promise",
                        hook=(
                            "A child and an otter guardian carry drifting lanterns across the "
                            "harbor so each light can return home before bed."
                        ),
                        central_conflict=(
                            "They need to help each lantern find the right doorstep before one "
                            "worried pause stretches the journey too long."
                        ),
                        why_it_fits=(
                            "It fits the selected lane with soft movement, harbor wonder, and "
                            "a clearly restful ending."
                        ),
                    ),
                    GeneratedPitchCandidate(
                        title="The Last Lantern Question",
                        hook=(
                            "When one lantern stays awake after the docks quiet down, the child "
                            "follows it to learn what the night still needs."
                        ),
                        central_conflict=(
                            "The answer must be found before everyone can settle, but each "
                            "discovery has to stay gentle and reassuring."
                        ),
                        why_it_fits=(
                            "It fits the brief by turning the harbor imagery into a bedtime-safe "
                            "mystery with quick emotional repair."
                        ),
                    ),
                    GeneratedPitchCandidate(
                        title="Moonpost Harbor Map",
                        hook=(
                            "A hidden path of lights appears across the docks, leading the child "
                            "and otter toward the night's unfinished kindness."
                        ),
                        central_conflict=(
                            "Each stop asks them to choose patience and care before the last "
                            "route home can become real."
                        ),
                        why_it_fits=(
                            "It fits Quest Fantasy and Hushed Wonder by keeping the stakes soft "
                            "while the imagery stays luminous."
                        ),
                    ),
                ]
            ),
            raw_response={
                "stub": True,
                "usageMetadata": {
                    "promptTokenCount": 540,
                    "candidatesTokenCount": 210,
                    "totalTokenCount": 750,
                },
            },
        )

    def close(self) -> None:
        return None


class StubCharacterGenerationAdapter:
    def __init__(self) -> None:
        self.model_id = "gemini-3.1-pro"

    def generate(
        self,
        invocation: CharacterGenerationInvocation,
    ) -> CharacterGenerationInvocationResult:
        return CharacterGenerationInvocationResult(
            invocation=invocation,
            structured_output=CharacterGenerationStructuredOutput(
                character_sheets=[
                    GeneratedCharacterSheetCandidate(
                        title="Juniper Keeper Cast",
                        summary=(
                            "Mira leads the harbor story with a calm helper dynamic and a "
                            "clear bedtime-safe repair arc."
                        ),
                        story_function=(
                            "The cast supports the pitch by making later beats easy to outline "
                            "around visible care and emotional repair."
                        ),
                        bedtime_safety_notes=(
                            "Every worry is buffered by visible helpers, named feelings, and a "
                            "clearly restorative ending."
                        ),
                        visual_motifs=["lantern glow", "moonlit docks", "otter paws"],
                        protagonist=GeneratedCharacterProfile(
                            name="Mira",
                            role="sleepy lantern-keeper in training",
                            goal="guide the last harbor lights home before everyone settles",
                            flaw="tries to fix every worry alone before asking for help",
                            comfort_trait="counts steady reflections until breathing slows",
                            bedtime_safety_notes=(
                                "Mira stays emotionally safe because helpers remain close and "
                                "calm in every scene."
                            ),
                            relationships=[
                                "Trusts Otis to steady the plan when emotions spike.",
                            ],
                            visual_anchors=["lantern sleeves", "soft satchel"],
                        ),
                        supporting_cast=[
                            GeneratedCharacterProfile(
                                name="Otis",
                                role="patient otter guardian",
                                goal="help Mira slow the pacing whenever the night feels bigger",
                                flaw="over-explains instead of letting Mira discover the answer",
                                comfort_trait="grounds scenes with practical rituals",
                                bedtime_safety_notes=(
                                    "Otis makes each obstacle small, readable, and quickly "
                                    "reassuring."
                                ),
                                relationships=[
                                    "Acts as Mira's calm sounding board.",
                                ],
                                visual_anchors=["river coat", "tidy satchel"],
                            )
                        ],
                    ),
                    GeneratedCharacterSheetCandidate(
                        title="Harbor Listener Cast",
                        summary=(
                            "Pip centers the same pitch around noticing, conversation, and a "
                            "gentler mystery rhythm."
                        ),
                        story_function=(
                            "The cast makes the pitch emotionally observant so later beats can "
                            "lean on dialogue and calm discovery."
                        ),
                        bedtime_safety_notes=(
                            "The mystery stays bedtime-safe because each reveal is named, "
                            "shared, and softened right away."
                        ),
                        visual_motifs=["still water", "sleepy reeds", "silver light"],
                        protagonist=GeneratedCharacterProfile(
                            name="Pip",
                            role="quiet listener for the sleeping shoreline",
                            goal="understand what the night still needs so everyone can settle",
                            flaw="hesitates to speak when a worried feeling first appears",
                            comfort_trait="matches footsteps to the hush of water",
                            bedtime_safety_notes=(
                                "Pip's worries stay readable and quickly supported by the whole "
                                "cast."
                            ),
                            relationships=[
                                "Leans on Rowan when emotions need clearer words.",
                            ],
                            visual_anchors=["silver pockets", "reed-shadow patterns"],
                        ),
                        supporting_cast=[
                            GeneratedCharacterProfile(
                                name="Rowan",
                                role="older sibling translator",
                                goal="turn worry into practical kindness whenever the mood dips",
                                flaw="tries to protect everyone before listening fully",
                                comfort_trait="restates feelings in calmer words",
                                bedtime_safety_notes=(
                                    "Rowan keeps discoveries gentle by naming them before they "
                                    "can feel overwhelming."
                                ),
                                relationships=[
                                    "Helps Pip speak the hidden worry aloud.",
                                ],
                                visual_anchors=["striped scarf", "reed basket"],
                            )
                        ],
                    ),
                    GeneratedCharacterSheetCandidate(
                        title="Moonpath Guide Cast",
                        summary=(
                            "Tavi shapes the harbor story around ritual, routes, and a compact "
                            "support team."
                        ),
                        story_function=(
                            "The cast supports the pitch with steady movement and small repairs "
                            "that stay easy to stage later."
                        ),
                        bedtime_safety_notes=(
                            "The route home stays emotionally safe because every stop has a "
                            "helper and a clearly calming purpose."
                        ),
                        visual_motifs=["mapped lights", "quiet ropes", "silver water"],
                        protagonist=GeneratedCharacterProfile(
                            name="Tavi",
                            role="gentle route-maker for sleepy travelers",
                            goal="lead one lingering lantern back toward a safe ending",
                            flaw="overplans each step and worries when the route changes",
                            comfort_trait=(
                                "returns to familiar rituals whenever the story needs grounding"
                            ),
                            bedtime_safety_notes=(
                                "Tavi's planning worries stay manageable because the support cast "
                                "keeps the route readable and warm."
                            ),
                            relationships=[
                                "Relies on Ember to reframe obstacles as invitations.",
                            ],
                            visual_anchors=["mapped ribbon", "quiet satchel"],
                        ),
                        supporting_cast=[
                            GeneratedCharacterProfile(
                                name="Ember",
                                role="warmhearted mentor",
                                goal=(
                                    "keep the quest readable and calm without flattening the wonder"
                                ),
                                flaw="tries to solve problems too early",
                                comfort_trait="grounds scenes with a steady ritual phrase",
                                bedtime_safety_notes=(
                                    "Ember gives every obstacle a gentle explanation before the "
                                    "story moves on."
                                ),
                                relationships=[
                                    "Encourages Tavi to trust the quieter route.",
                                ],
                                visual_anchors=["soft lantern staff", "threaded map"],
                            )
                        ],
                    ),
                ]
            ),
            raw_response={
                "stub": True,
                "usageMetadata": {
                    "promptTokenCount": 640,
                    "candidatesTokenCount": 260,
                    "totalTokenCount": 900,
                },
            },
        )

    def close(self) -> None:
        return None


class StubBeatSheetGenerationAdapter:
    def __init__(self) -> None:
        self.model_id = "gemini-3.1-pro"

    def generate(
        self,
        invocation: BeatSheetGenerationInvocation,
    ) -> BeatSheetGenerationInvocationResult:
        return BeatSheetGenerationInvocationResult(
            invocation=invocation,
            structured_output=BeatSheetGenerationStructuredOutput(
                beat_sheet=GeneratedBeatSheetCandidate(
                    summary=(
                        "A harbor Save-the-Cat arc where Mira's wonder turns into a restful "
                        "sense of belonging."
                    ),
                    bedtime_notes=(
                        "Each pressure beat is kept soft, visible, and quickly buffered by "
                        "companionship."
                    ),
                    beats=[
                        GeneratedBeatSheetBeat(
                            key=key,
                            label=label,
                            summary=(
                                f"{label} keeps the harbor story high level, concrete, and "
                                "bedtime-ready."
                            ),
                            emotional_intent=(
                                f"{label} should feel calm, clear, and emotionally readable."
                            ),
                            bedtime_softening_note=(
                                f"{label} stays bedtime-safe by keeping tension gentle and "
                                "quickly soothed."
                            ),
                        )
                        for key, label in (
                            ("opening_image", "Opening Image"),
                            ("theme_stated", "Theme Stated"),
                            ("set_up", "Set-Up"),
                            ("catalyst", "Catalyst"),
                            ("debate", "Debate"),
                            ("break_into_two", "Break into Two"),
                            ("b_story", "B Story"),
                            ("fun_and_games", "Fun and Games"),
                            ("midpoint", "Midpoint"),
                            ("bad_guys_close_in", "Bad Guys Close In"),
                            ("all_is_lost", "All Is Lost"),
                            ("dark_night_of_the_soul", "Dark Night of the Soul"),
                            ("break_into_three", "Break into Three"),
                            ("finale", "Finale"),
                            ("final_image", "Final Image"),
                        )
                    ],
                )
            ),
            raw_response={
                "stub": True,
                "usageMetadata": {
                    "promptTokenCount": 820,
                    "candidatesTokenCount": 340,
                    "totalTokenCount": 1160,
                },
            },
        )

    def close(self) -> None:
        return None


@pytest.fixture
def session_api_client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Iterator[TestClient]:
    database_path = tmp_path / "session-api.sqlite3"
    monkeypatch.setenv("STORYTELLER_DATABASE_URL", f"sqlite+pysqlite:///{database_path}")

    get_settings.cache_clear()
    get_engine.cache_clear()
    get_session_factory.cache_clear()

    engine = get_engine()
    Base.metadata.create_all(engine)

    fake_storage = InMemoryObjectStorage(get_settings())
    monkeypatch.setattr("app.main.build_object_storage_service", lambda settings: fake_storage)

    app = create_app()
    app.state.beat_sheet_generation_adapter = StubBeatSheetGenerationAdapter()
    app.state.brief_normalization_adapter = StubBriefNormalizationAdapter()
    app.state.character_generation_adapter = StubCharacterGenerationAdapter()
    app.state.pitch_generation_adapter = StubPitchGenerationAdapter()

    with TestClient(app) as test_client:
        yield test_client

    get_settings.cache_clear()
    get_engine.cache_clear()
    get_session_factory.cache_clear()


def _seed_catalog_rows() -> None:
    db_session = get_session_factory()()
    try:
        seed_catalog(db_session, load_catalog_document(CATALOG_FILE_PATH))
    finally:
        db_session.close()


def _create_composition_ready_session_via_api(
    session_api_client: TestClient,
) -> dict[str, object]:
    _seed_catalog_rows()

    created = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Moonlit Harbor"},
    ).json()
    session_id = created["id"]

    assert (
        session_api_client.post(
            f"/api/v1/sessions/{session_id}/selections/genre",
            json={"genre_slug": "quest-fantasy", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{session_id}/selections/tone",
            json={"tone_profile_slug": "hushed-wonder", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{session_id}/story-brief",
            json={
                "raw_brief": (
                    "A harbor fox follows a drifting bell across moonlit water and brings the "
                    "night back to calm."
                ),
                "origin": "workspace",
            },
        ).status_code
        == 200
    )

    pitch_payload = session_api_client.post(
        f"/api/v1/sessions/{session_id}/pitches/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    first_pitch_id = pitch_payload["snapshot"]["pitch_batches"][0]["pitches"][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{session_id}/selections/pitch",
            json={"pitch_id": first_pitch_id, "origin": "workspace"},
        ).status_code
        == 200
    )

    character_payload = session_api_client.post(
        f"/api/v1/sessions/{session_id}/characters/generate",
        json={"candidate_count": 2, "origin": "workspace"},
    ).json()
    first_character_id = character_payload["snapshot"]["character_sheet_batches"][0][
        "character_sheets"
    ][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{session_id}/selections/character-sheet",
            json={"character_sheet_id": first_character_id, "origin": "workspace"},
        ).status_code
        == 200
    )

    beat_payload = session_api_client.post(
        f"/api/v1/sessions/{session_id}/beats/generate",
        json={"origin": "workspace"},
    ).json()
    first_beat_id = beat_payload["snapshot"]["beat_sheet_revisions"][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{session_id}/selections/beat-sheet",
            json={"beat_sheet_id": first_beat_id, "origin": "workspace"},
        ).status_code
        == 200
    )

    story_setup_payload = session_api_client.post(
        f"/api/v1/sessions/{session_id}/story-setup",
        json={
            "target_word_count": 1800,
            "target_runtime_minutes": 12,
            "chapter_count": 3,
            "approximate_scene_count": 8,
            "guidance_notes": "Let each chapter land calmer than it began.",
            "origin": "workspace",
        },
    ).json()

    return {
        "session_id": session_id,
        "snapshot": story_setup_payload["snapshot"],
    }


def test_list_recent_sessions_endpoint_returns_sessions_with_latest_first(
    session_api_client: TestClient,
) -> None:
    db_session = get_session_factory()()
    try:
        service = SessionService(db_session)
        older = service.create_session(working_title="Older Session")
        newer = service.create_session(working_title="Newer Session")

        older_row = db_session.get(StorySession, older.id)
        newer_row = db_session.get(StorySession, newer.id)
        assert older_row is not None and newer_row is not None

        older_row.updated_at = datetime.now(timezone.utc) - timedelta(days=1)
        newer_row.updated_at = datetime.now(timezone.utc)
        db_session.commit()

        service.update_stage_state(
            newer.id,
            stage=WorkflowStage.GENRE,
            status=WorkflowStageState.COMPLETED,
            detail="Accepted quest fantasy.",
        )
    finally:
        db_session.close()

    response = session_api_client.get("/api/v1/sessions")

    assert response.status_code == 200
    payload = response.json()

    assert [session["display_title"] for session in payload[:2]] == [
        "Newer Session",
        "Older Session",
    ]
    assert payload[0]["overall_status"] == "in_progress"
    assert payload[0]["current_stage"] == "tone"
    assert payload[0]["progress"]["completed_stages"] == 1
    assert payload[1]["overall_status"] == "draft"
    assert payload[1]["progress"]["completed_stages"] == 0


def test_session_usage_endpoint_returns_rollups_and_recent_calls(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post("/api/v1/sessions", json={"working_title": "Usage"})
    assert create_response.status_code == 201
    session_id = create_response.json()["id"]

    db_session = get_session_factory()()
    try:
        usage = SessionModelUsageService(db_session)
        usage.record_model_call(
            context=ModelUsageContext(
                session_id=session_id,
                usage_bucket=ModelUsageBucket.PLANNING,
                workflow_stage=WorkflowStage.BRIEF,
                purpose="brief_normalization",
                model_id="gemini-3.1-flash-lite",
                prompt_version="brief_normalizer.v1",
            ),
            elapsed_ms=420,
            outcome=ModelCallOutcome.SUCCEEDED,
            raw_response={
                "usageMetadata": {
                    "promptTokenCount": 240,
                    "candidatesTokenCount": 72,
                    "totalTokenCount": 312,
                }
            },
        )
        usage.record_model_call(
            context=ModelUsageContext(
                session_id=session_id,
                usage_bucket=ModelUsageBucket.PLANNING,
                workflow_stage=WorkflowStage.PITCHES,
                purpose="pitch_generation",
                model_id="gemini-3.1-pro",
                prompt_version="pitch_generation.v3",
            ),
            elapsed_ms=860,
            outcome=ModelCallOutcome.SUCCEEDED,
            raw_response={
                "usageMetadata": {
                    "promptTokenCount": 540,
                    "candidatesTokenCount": 210,
                    "totalTokenCount": 750,
                }
            },
        )
        db_session.commit()
    finally:
        db_session.close()

    usage_response = session_api_client.get(f"/api/v1/sessions/{session_id}/usage")
    assert usage_response.status_code == 200
    payload = usage_response.json()

    planning_bucket = next(
        bucket for bucket in payload["summary"]["buckets"] if bucket["usage_bucket"] == "planning"
    )
    assert planning_bucket["total_calls"] == 2
    assert planning_bucket["token_usage"]["total_tokens"] == 1062
    assert planning_bucket["approximate_cost_usd"] is not None
    assert planning_bucket["models_used"] == ["gemini-3.1-flash-lite", "gemini-3.1-pro"]
    assert payload["recent_calls"][0]["purpose"] == "pitch_generation"
    assert payload["recent_calls"][1]["purpose"] == "brief_normalization"


def test_get_session_snapshot_endpoint_returns_full_snapshot(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Moonlit Harbor"},
    )
    created = create_response.json()

    response = session_api_client.get(f"/api/v1/sessions/{created['id']}")

    assert response.status_code == 200
    payload = response.json()

    assert payload["id"] == created["id"]
    assert payload["display_title"] == "Moonlit Harbor"
    assert payload["current_stage"] == "genre"
    assert payload["progress"]["total_stages"] == 10
    assert len(payload["stage_states"]) == 10
    assert payload["stage_states"][0]["stage"] == "genre"
    assert payload["stage_states"][0]["status"] == "draft"
    assert payload["agent_context_summary"].startswith("Session title: Moonlit Harbor")


def test_get_genre_catalog_endpoint_returns_seeded_genres(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()

    response = session_api_client.get("/api/v1/catalog/genres")

    assert response.status_code == 200
    payload = response.json()

    assert [genre["slug"] for genre in payload] == [
        "quest-fantasy",
        "gentle-mystery",
        "cozy-animal-adventure",
        "magical-friendship",
        "dreamworld-voyage",
        "soft-science-wonder",
    ]
    assert payload[0]["bedtime_safety_notes"].startswith("Keep the journey wondrous")
    assert payload[0]["arc_notes"]["core_arc"].startswith("Leave a familiar safe place")


def test_get_tone_catalog_endpoint_returns_only_requested_genre_tones(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()

    response = session_api_client.get("/api/v1/catalog/genres/quest-fantasy/tones")

    assert response.status_code == 200
    payload = response.json()

    assert [tone["slug"] for tone in payload] == [
        "hushed-wonder",
        "lantern-brave",
        "fireside-reassurance",
    ]
    assert payload[0]["description"].startswith("Moonlit awe")
    assert payload[0]["descriptors"] == [
        "luminous",
        "moonlit",
        "gentle",
        "reverent",
    ]
    assert payload[0]["default_planning_hints"]["ending_style"].startswith("return home carrying")


def test_get_tone_catalog_endpoint_returns_404_for_unknown_genre(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()

    response = session_api_client.get("/api/v1/catalog/genres/unknown-genre/tones")

    assert response.status_code == 404
    assert response.json()["detail"] == "no active genre matched 'unknown-genre'"


def test_hydrate_session_endpoint_returns_snapshot_history_and_metadata(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Moonlit Harbor"},
    )
    created = create_response.json()

    response = session_api_client.get(f"/api/v1/sessions/{created['id']}/hydrate")

    assert response.status_code == 200
    payload = response.json()

    assert payload["snapshot"]["id"] == created["id"]
    assert payload["recent_history"]["session_id"] == created["id"]
    assert payload["recent_history"]["events"][0]["event_type"] == "session.created"
    assert payload["hydration"]["strategy"] == "materialized_only"
    assert payload["hydration"]["latest_sequence_number"] == 1
    assert payload["hydration"]["history_event_count"] == 1


def test_hydrate_session_endpoint_includes_audio_segments_and_preview_urls(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Moonlit Harbor"},
    )
    created = create_response.json()

    db_session = get_session_factory()()
    try:
        audio_job = AudioJob(
            session_id=created["id"],
            status=JobStatus.IN_PROGRESS,
            voice_key="moonbeam",
            playback_speed=0.95,
            include_background_music=False,
            music_profile="lullaby_piano",
            estimated_duration_seconds=780,
            current_segment_index=2,
            config_json={
                "total_segments": 2,
                "completed_segments": 1,
                "progress_percent": 54.0,
                "current_step": "Rendering narration segment 2 of 2.",
                "current_step_index": 2,
                "total_steps": 4,
            },
        )
        db_session.add(audio_job)
        db_session.flush()

        db_session.add_all(
            [
                NarrationSegment(
                    session_id=created["id"],
                    audio_job_id=audio_job.id,
                    segment_index=1,
                    status=JobStatus.COMPLETED,
                    source_boundary_kind=NarrationSourceBoundaryKind.CHAPTER,
                    source_outline_card_title="Lantern launch",
                    text_content=(
                        "Mira set the first lantern on the water and waited for "
                        "the harbor to answer."
                    ),
                    word_count=14,
                    text_start_offset=0,
                    text_end_offset=76,
                    pause_after_seconds=3,
                    pause_hint=NarrationPauseHint.CHAPTER_BREAK,
                    metadata_json={"split_reason": "paragraph_boundary"},
                ),
                NarrationSegment(
                    session_id=created["id"],
                    audio_job_id=audio_job.id,
                    segment_index=2,
                    status=JobStatus.QUEUED,
                    source_boundary_kind=NarrationSourceBoundaryKind.CHAPTER,
                    source_outline_card_title="Silver bell crossing",
                    text_content=(
                        "Otis stayed close while the silver bell called from the cove."
                    ),
                    word_count=11,
                    text_start_offset=78,
                    text_end_offset=140,
                    pause_after_seconds=0,
                    pause_hint=NarrationPauseHint.NONE,
                    metadata_json={"split_reason": "sentence_boundary"},
                ),
            ]
        )
        db_session.flush()

        db_session.add(
            SessionAsset(
                session_id=created["id"],
                audio_job_id=audio_job.id,
                asset_kind=AssetKind.AUDIO_SEGMENT,
                status=AssetStatus.READY,
                storage_bucket="storyteller-audio",
                object_path=(
                    f"sessions/{created['id']}/audio/jobs/{audio_job.id}/segments/0001.wav"
                ),
                mime_type="audio/wav",
                segment_index=1,
            )
        )
        db_session.commit()
    finally:
        db_session.close()

    response = session_api_client.get(f"/api/v1/sessions/{created['id']}/hydrate")

    assert response.status_code == 200
    payload = response.json()

    assert len(payload["snapshot"]["audio_segments"]) == 2
    assert payload["snapshot"]["audio_segments"][0]["source_outline_card_title"] == (
        "Lantern launch"
    )
    assert (
        payload["snapshot"]["audio_segments"][0]["preview_asset"]["access"]["stream_path"]
        == f"/api/v1/sessions/{created['id']}/assets/"
        f"{payload['snapshot']['audio_segments'][0]['preview_asset']['id']}/content"
        "?disposition=inline"
    )
    assert payload["snapshot"]["audio_segments"][1]["split_reason"] == "sentence_boundary"


def test_get_session_asset_content_supports_byte_ranges_for_audio(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Audio Asset Range"},
    )
    created = create_response.json()
    fake_storage = session_api_client.app.state.object_storage
    assert isinstance(fake_storage, InMemoryObjectStorage)
    audio_location = fake_storage.paths.final_audio(
        session_id=created["id"],
        job_id="audio-1",
    )
    audio_bytes = b"FAKEAUDIO"
    fake_storage.upload_bytes(audio_location, audio_bytes, content_type="audio/mpeg")

    db_session = get_session_factory()()
    try:
        audio_asset = SessionAsset(
            session_id=created["id"],
            asset_kind=AssetKind.FINAL_AUDIO,
            status=AssetStatus.READY,
            storage_bucket=audio_location.bucket,
            object_path=audio_location.key,
            mime_type="audio/mpeg",
        )
        db_session.add(audio_asset)
        db_session.commit()
        asset_id = audio_asset.id
    finally:
        db_session.close()

    response = session_api_client.get(
        f"/api/v1/sessions/{created['id']}/assets/{asset_id}/content",
        headers={"Range": "bytes=0-3"},
    )

    assert response.status_code == 206
    assert response.content == b"FAKE"
    assert response.headers["content-range"] == f"bytes 0-3/{len(audio_bytes)}"
    assert response.headers["accept-ranges"] == "bytes"
    assert response.headers["content-type"] == "audio/mpeg"


def test_get_named_story_docx_artifact_generates_export_from_story_text(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Lantern Harbor"},
    )
    created = create_response.json()
    fake_storage = session_api_client.app.state.object_storage
    assert isinstance(fake_storage, InMemoryObjectStorage)

    story_location = fake_storage.paths.export_asset(
        session_id=created["id"],
        export_kind="story",
        export_id="accepted-manuscript",
        extension="md",
    )
    fake_storage.upload_text(
        story_location,
        "# Chapter 1\n\nMira carried the lantern home.",
        content_type="text/markdown; charset=utf-8",
    )

    db_session = get_session_factory()()
    try:
        story_asset = SessionAsset(
            session_id=created["id"],
            asset_kind=AssetKind.STORY_TEXT,
            status=AssetStatus.READY,
            storage_bucket=story_location.bucket,
            object_path=story_location.key,
            mime_type="text/markdown",
        )
        db_session.add(story_asset)
        db_session.commit()
        source_story_asset_id = story_asset.id
    finally:
        db_session.close()

    response = session_api_client.get(
        f"/api/v1/sessions/{created['id']}/artifacts/story-docx",
    )

    assert response.status_code == 200
    assert (
        response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment" in response.headers["content-disposition"]

    document = Document(BytesIO(response.content))
    assert document.paragraphs[0].text == "Lantern Harbor"
    assert any(
        paragraph.text == "Mira carried the lantern home."
        for paragraph in document.paragraphs
    )

    db_session = get_session_factory()()
    try:
        exported_asset = db_session.execute(
            select(SessionAsset).where(
                SessionAsset.session_id == created["id"],
                SessionAsset.asset_kind == AssetKind.STORY_DOCX,
            )
        ).scalar_one()
        assert exported_asset.status == AssetStatus.READY
        assert exported_asset.metadata_json["source_story_asset_id"] == source_story_asset_id
        exported_bytes = fake_storage.download_bytes(
            fake_storage.paths.export_asset(
                session_id=created["id"],
                export_kind="docx",
                export_id="final-manuscript",
                extension="docx",
            )
        )
        assert exported_bytes == response.content
    finally:
        db_session.close()


def test_post_story_docx_artifact_generates_export_metadata(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Lantern Harbor"},
    )
    created = create_response.json()
    fake_storage = session_api_client.app.state.object_storage
    assert isinstance(fake_storage, InMemoryObjectStorage)

    story_location = fake_storage.paths.export_asset(
        session_id=created["id"],
        export_kind="story",
        export_id="accepted-manuscript",
        extension="md",
    )
    fake_storage.upload_text(
        story_location,
        "# Chapter 1\n\nMira carried the lantern home.",
        content_type="text/markdown; charset=utf-8",
    )

    db_session = get_session_factory()()
    try:
        db_session.add(
            SessionAsset(
                session_id=created["id"],
                asset_kind=AssetKind.STORY_TEXT,
                status=AssetStatus.READY,
                storage_bucket=story_location.bucket,
                object_path=story_location.key,
                mime_type="text/markdown",
            )
        )
        db_session.commit()
    finally:
        db_session.close()

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/artifacts/story-docx",
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["asset_kind"] == "story_docx"
    assert payload["status"] == "ready"
    assert payload["access"]["filename"] == "Lantern-Harbor.docx"
    assert payload["access"]["download_path"].endswith(
        f"/api/v1/sessions/{created['id']}/assets/{payload['id']}/content?disposition=attachment"
    )

    db_session = get_session_factory()()
    try:
        exported_asset = db_session.get(SessionAsset, payload["id"])
        assert exported_asset is not None
        assert exported_asset.asset_kind == AssetKind.STORY_DOCX
        assert exported_asset.status == AssetStatus.READY
    finally:
        db_session.close()


def test_get_session_history_endpoint_returns_durable_timeline(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "History Check"},
    )
    created = create_response.json()

    response = session_api_client.get(f"/api/v1/sessions/{created['id']}/history")

    assert response.status_code == 200
    payload = response.json()

    assert payload["session_id"] == created["id"]
    assert payload["latest_sequence_number"] == 1
    assert payload["events"][0]["event_type"] == "session.created"


def test_record_session_ui_action_endpoint_returns_recorded_event(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Echo Source"},
    )
    created = create_response.json()

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/ui-actions",
        json={
            "action": "navigate_to_stage",
            "stage": "audio",
            "control_id": "stage-navigator",
            "value_summary": "Audio",
            "origin": "workspace",
        },
    )

    assert response.status_code == 201
    payload = response.json()

    assert payload["event_type"] == "ui.action.recorded"
    assert payload["stage"] == "audio"
    assert payload["payload"]["action"] == "navigate_to_stage"
    assert payload["payload"]["control_id"] == "stage-navigator"
    assert payload["payload"]["value_summary"] == "Audio"
    assert payload["payload"]["origin"] == "workspace"


def test_select_session_genre_endpoint_persists_choice_and_returns_snapshot(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Genre API"},
    )
    created = create_response.json()

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/selections/genre",
        json={
            "genre_slug": "quest-fantasy",
            "origin": "workspace",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "selection.recorded"
    assert payload["event"]["stage"] == "genre"
    assert payload["event"]["payload"]["selection_kind"] == "genre"
    assert payload["snapshot"]["selected_genre"]["slug"] == "quest-fantasy"
    assert payload["snapshot"]["current_stage"] == "tone"
    assert payload["snapshot"]["resume_stage"] == "tone"
    assert payload["snapshot"]["stage_states"][0]["status"] == "completed"
    assert payload["snapshot"]["stage_states"][0]["last_event_summary"] == (
        "Selected genre: Quest Fantasy."
    )


def test_select_session_tone_endpoint_persists_choice_and_returns_snapshot(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Tone API"},
    )
    created = create_response.json()

    genre_response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/selections/genre",
        json={
            "genre_slug": "quest-fantasy",
            "origin": "workspace",
        },
    )
    assert genre_response.status_code == 200

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/selections/tone",
        json={
            "tone_profile_slug": "hushed-wonder",
            "origin": "workspace",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "selection.recorded"
    assert payload["event"]["stage"] == "tone"
    assert payload["event"]["payload"]["selection_kind"] == "tone_profile"
    assert payload["snapshot"]["selected_tone_profile"]["slug"] == "hushed-wonder"
    assert payload["snapshot"]["current_stage"] == "brief"
    assert payload["snapshot"]["resume_stage"] == "brief"
    assert payload["snapshot"]["stage_states"][1]["status"] == "completed"
    assert payload["snapshot"]["stage_states"][1]["last_event_summary"] == (
        "Selected tone profile: Hushed Wonder."
    )


def test_select_session_tone_endpoint_requires_genre_first(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Tone API Guardrail"},
    )
    created = create_response.json()

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/selections/tone",
        json={
            "tone_profile_slug": "hushed-wonder",
            "origin": "workspace",
        },
    )

    assert response.status_code == 422
    assert response.json()["detail"] == "select a genre before choosing a tone"


def test_save_story_brief_endpoint_persists_revisioned_brief_and_returns_snapshot(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Brief API"},
    )
    created = create_response.json()

    genre_response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/selections/genre",
        json={
            "genre_slug": "quest-fantasy",
            "origin": "workspace",
        },
    )
    assert genre_response.status_code == 200

    tone_response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/selections/tone",
        json={
            "tone_profile_slug": "hushed-wonder",
            "origin": "workspace",
        },
    )
    assert tone_response.status_code == 200

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/story-brief",
        json={
            "story_idea": (
                "A child follows floating lanterns through a harbor and helps a shy "
                "otter guide each light back to bed."
            ),
            "desired_themes": "gentle bravery, belonging, quiet wonder",
            "key_images": "floating lanterns, moonlit water, otter paws on a skiff rail",
            "audience_notes": "Ideal for kids who like wonder without spooky twists.",
            "must_have_elements": "End with the harbor settled and the child tucked in safely.",
            "origin": "workspace",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "content.user_edit.recorded"
    assert payload["event"]["stage"] == "brief"
    assert payload["event"]["payload"]["target_kind"] == "story_brief"
    assert payload["event"]["payload"]["revision_number"] == 1
    assert payload["snapshot"]["current_stage"] == "pitches"
    assert payload["snapshot"]["resume_stage"] == "pitches"
    assert payload["snapshot"]["story_brief"]["story_idea"].startswith(
        "A child follows floating lanterns"
    )
    assert (
        payload["snapshot"]["story_brief"]["normalized_summary"]
        == "A harbor bedtime quest where each lantern return settles the night."
    )
    assert (
        payload["snapshot"]["story_brief"]["normalized_preferences"]["setting"]
        == "a moonlit harbor"
    )
    assert (
        "Desired themes: gentle bravery, belonging, quiet wonder"
        in payload["snapshot"]["story_brief"]["raw_brief"]
    )
    assert payload["snapshot"]["stage_states"][2]["status"] == "completed"
    assert payload["snapshot"]["stage_states"][2]["detail"].startswith("Saved story brief:")


def test_save_story_brief_endpoint_requires_tone_prerequisite(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Brief API Guardrail"},
    )
    created = create_response.json()

    genre_response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/selections/genre",
        json={
            "genre_slug": "quest-fantasy",
            "origin": "workspace",
        },
    )
    assert genre_response.status_code == 200

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/story-brief",
        json={
            "story_idea": "A soft harbor lantern story.",
            "origin": "workspace",
        },
    )

    assert response.status_code == 409
    assert "cannot set 'brief' to 'completed'" in response.json()["detail"]


def test_generate_session_pitches_endpoint_persists_a_durable_batch(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    created = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Pitch API"},
    ).json()

    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/genre",
            json={"genre_slug": "quest-fantasy", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/tone",
            json={"tone_profile_slug": "hushed-wonder", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/story-brief",
            json={
                "story_idea": (
                    "A child follows floating lanterns through a harbor and helps a shy otter "
                    "guide each light back to bed."
                ),
                "origin": "workspace",
            },
        ).status_code
        == 200
    )

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/pitches/generate",
        json={
            "candidate_count": 3,
            "guidance": "Keep the harbor mystery very gentle.",
            "origin": "workspace",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "ai.output.recorded"
    assert payload["event"]["payload"]["output_kind"] == "pitch_batch"
    assert payload["event"]["payload"]["candidate_count"] == 3
    assert payload["snapshot"]["current_stage"] == "pitches"
    assert payload["snapshot"]["resume_stage"] == "pitches"
    assert payload["snapshot"]["selected_pitch"] is None
    assert payload["snapshot"]["pitch_batches"][0]["candidate_count"] == 3
    assert payload["snapshot"]["pitch_batches"][0]["pitches"][0]["title"] == (
        "The Juniper Lake Promise"
    )
    assert payload["snapshot"]["stage_states"][3]["status"] == "in_progress"
    assert payload["snapshot"]["stage_states"][3]["detail"] == (
        "Generated 3 pitch options. Select one to continue."
    )


def test_select_session_pitch_endpoint_marks_choice_and_advances_to_characters(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    created = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Pitch Selection API"},
    ).json()

    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/genre",
            json={"genre_slug": "quest-fantasy", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/tone",
            json={"tone_profile_slug": "hushed-wonder", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/story-brief",
            json={
                "story_idea": (
                    "A child follows floating lanterns through a harbor and helps a shy otter "
                    "guide each light back to bed."
                ),
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    generated = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/pitches/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    first_pitch_id = generated["snapshot"]["pitch_batches"][0]["pitches"][0]["id"]

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/selections/pitch",
        json={
            "pitch_id": first_pitch_id,
            "origin": "workspace",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "selection.recorded"
    assert payload["event"]["payload"]["selection_kind"] == "pitch"
    assert payload["snapshot"]["current_stage"] == "characters"
    assert payload["snapshot"]["resume_stage"] == "characters"
    assert payload["snapshot"]["selected_pitch"]["id"] == first_pitch_id
    assert payload["snapshot"]["selected_pitch"]["is_selected"] is True
    assert payload["snapshot"]["stage_states"][3]["status"] == "completed"


def test_refine_session_pitch_endpoint_creates_a_selected_refined_pitch(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    created = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Pitch Refinement API"},
    ).json()

    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/genre",
            json={"genre_slug": "quest-fantasy", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/tone",
            json={"tone_profile_slug": "hushed-wonder", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/story-brief",
            json={
                "story_idea": (
                    "A child follows floating lanterns through a harbor and helps a shy otter "
                    "guardian guide each light back to bed."
                ),
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    generated = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/pitches/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    source_pitch = generated["snapshot"]["pitch_batches"][0]["pitches"][1]

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/pitches/refine",
        json={
            "pitch_id": source_pitch["id"],
            "instructions": "Make it about siblings who help each other settle down.",
            "origin": "chat",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "selection.recorded"
    assert payload["event"]["payload"]["selection_kind"] == "pitch"
    assert "Refined from" in payload["event"]["payload"]["rationale"]
    assert payload["snapshot"]["current_stage"] == "characters"
    assert payload["snapshot"]["selected_pitch"]["generation_kind"] == "refinement"
    assert payload["snapshot"]["selected_pitch"]["source_pitch_id"] == source_pitch["id"]
    assert payload["snapshot"]["selected_pitch"]["refinement_instructions"] == (
        "Make it about siblings who help each other settle down."
    )
    assert payload["snapshot"]["pitch_batches"][0]["generation_kind"] == "refinement"
    assert payload["snapshot"]["pitch_batches"][0]["source_pitch_title"] == source_pitch["title"]


def test_generate_session_character_sheets_endpoint_persists_a_durable_batch(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    created = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Character API"},
    ).json()

    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/genre",
            json={"genre_slug": "quest-fantasy", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/tone",
            json={"tone_profile_slug": "hushed-wonder", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/story-brief",
            json={
                "story_idea": (
                    "A child follows floating lanterns through a harbor and helps a shy otter "
                    "guardian guide each light back to bed."
                ),
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    generated_pitches = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/pitches/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    selected_pitch_id = generated_pitches["snapshot"]["pitch_batches"][0]["pitches"][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/pitch",
            json={"pitch_id": selected_pitch_id, "origin": "workspace"},
        ).status_code
        == 200
    )

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/characters/generate",
        json={
            "candidate_count": 3,
            "guidance": "Keep the support cast compact and clearly cozy.",
            "origin": "workspace",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "ai.output.recorded"
    assert payload["event"]["payload"]["output_kind"] == "character_sheet"
    assert payload["event"]["payload"]["candidate_count"] == 3
    assert payload["snapshot"]["current_stage"] == "characters"
    assert payload["snapshot"]["resume_stage"] == "characters"
    assert payload["snapshot"]["selected_character_sheet"] is None
    assert payload["snapshot"]["character_sheet_batches"][0]["candidate_count"] == 3
    assert payload["snapshot"]["character_sheet_batches"][0]["character_sheets"][0]["title"] == (
        "Juniper Keeper Cast"
    )
    assert payload["snapshot"]["stage_states"][4]["status"] == "in_progress"
    assert payload["snapshot"]["stage_states"][4]["detail"] == (
        "Generated 3 character options. Select one to continue."
    )


def test_select_session_character_sheet_endpoint_marks_choice_and_advances_to_beats(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    created = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Character Selection API"},
    ).json()

    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/genre",
            json={"genre_slug": "quest-fantasy", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/tone",
            json={"tone_profile_slug": "hushed-wonder", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/story-brief",
            json={
                "story_idea": (
                    "A child follows floating lanterns through a harbor and helps a shy otter "
                    "guardian guide each light back to bed."
                ),
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    generated_pitches = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/pitches/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    selected_pitch_id = generated_pitches["snapshot"]["pitch_batches"][0]["pitches"][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/pitch",
            json={"pitch_id": selected_pitch_id, "origin": "workspace"},
        ).status_code
        == 200
    )
    generated_characters = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/characters/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    first_character_sheet_id = generated_characters["snapshot"]["character_sheet_batches"][0][
        "character_sheets"
    ][0]["id"]

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/selections/character-sheet",
        json={
            "character_sheet_id": first_character_sheet_id,
            "origin": "workspace",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "selection.recorded"
    assert payload["event"]["payload"]["selection_kind"] == "character_sheet"
    assert payload["snapshot"]["current_stage"] == "beats"
    assert payload["snapshot"]["resume_stage"] == "beats"
    assert payload["snapshot"]["selected_character_sheet"]["id"] == first_character_sheet_id
    assert payload["snapshot"]["selected_character_sheet"]["is_selected"] is True
    assert payload["snapshot"]["stage_states"][4]["status"] == "completed"


def test_refine_session_character_sheet_endpoint_creates_a_selected_refined_sheet(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    created = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Character Refinement API"},
    ).json()

    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/genre",
            json={"genre_slug": "quest-fantasy", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/tone",
            json={"tone_profile_slug": "hushed-wonder", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/story-brief",
            json={
                "story_idea": (
                    "A child follows floating lanterns through a harbor and helps a shy otter "
                    "guardian guide each light back to bed."
                ),
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    generated_pitches = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/pitches/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    selected_pitch_id = generated_pitches["snapshot"]["pitch_batches"][0]["pitches"][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/pitch",
            json={"pitch_id": selected_pitch_id, "origin": "workspace"},
        ).status_code
        == 200
    )
    generated_characters = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/characters/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    source_character_sheet = generated_characters["snapshot"]["character_sheet_batches"][0][
        "character_sheets"
    ][1]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/character-sheet",
            json={
                "character_sheet_id": source_character_sheet["id"],
                "origin": "workspace",
            },
        ).status_code
        == 200
    )

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/characters/refine",
        json={
            "character_sheet_id": source_character_sheet["id"],
            "instructions": "Make the lead and guardian feel more sibling-like.",
            "focus_character_names": ["Pip", "Rowan"],
            "origin": "chat",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "selection.recorded"
    assert payload["event"]["payload"]["selection_kind"] == "character_sheet"
    assert "Refined from" in payload["event"]["payload"]["rationale"]
    assert payload["snapshot"]["current_stage"] == "beats"
    assert payload["snapshot"]["selected_character_sheet"]["generation_kind"] == "refinement"
    assert (
        payload["snapshot"]["selected_character_sheet"]["source_character_sheet_id"]
        == (source_character_sheet["id"])
    )
    assert payload["snapshot"]["selected_character_sheet"]["refinement_instructions"] == (
        "Make the lead and guardian feel more sibling-like."
    )
    assert payload["snapshot"]["character_sheet_batches"][0]["generation_kind"] == "refinement"
    assert (
        payload["snapshot"]["character_sheet_batches"][0]["source_character_sheet_title"]
        == source_character_sheet["title"]
    )


def test_generate_session_beat_sheet_endpoint_persists_a_revisioned_candidate(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    created = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Beat Sheet API"},
    ).json()

    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/genre",
            json={"genre_slug": "quest-fantasy", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/tone",
            json={"tone_profile_slug": "hushed-wonder", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/story-brief",
            json={
                "story_idea": (
                    "A child follows floating lanterns through a harbor and helps a shy otter "
                    "guardian guide each light back to bed."
                ),
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    generated_pitches = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/pitches/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    selected_pitch_id = generated_pitches["snapshot"]["pitch_batches"][0]["pitches"][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/pitch",
            json={"pitch_id": selected_pitch_id, "origin": "workspace"},
        ).status_code
        == 200
    )
    generated_characters = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/characters/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    selected_character_sheet_id = generated_characters["snapshot"]["character_sheet_batches"][0][
        "character_sheets"
    ][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/character-sheet",
            json={
                "character_sheet_id": selected_character_sheet_id,
                "origin": "workspace",
            },
        ).status_code
        == 200
    )

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/beats/generate",
        json={
            "guidance": "Keep the midpoint more awestruck than tense.",
            "focus_beats": ["midpoint"],
            "bedtime_goal": "land in a visibly sleepy ending",
            "origin": "workspace",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "ai.output.recorded"
    assert payload["event"]["payload"]["output_kind"] == "beat_sheet"
    assert payload["snapshot"]["current_stage"] == "beats"
    assert payload["snapshot"]["resume_stage"] == "beats"
    assert payload["snapshot"]["selected_beat_sheet"] is None
    assert payload["snapshot"]["beat_sheet_revisions"][0]["generation_kind"] == "generated"
    assert len(payload["snapshot"]["beat_sheet_revisions"][0]["beats"]) == 15
    assert payload["snapshot"]["stage_states"][5]["status"] == "in_progress"


def test_select_session_beat_sheet_endpoint_marks_choice_and_advances_to_story_setup(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    created = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Beat Selection API"},
    ).json()

    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/genre",
            json={"genre_slug": "quest-fantasy", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/tone",
            json={"tone_profile_slug": "hushed-wonder", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/story-brief",
            json={
                "story_idea": (
                    "A child follows floating lanterns through a harbor and helps a shy otter "
                    "guardian guide each light back to bed."
                ),
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    generated_pitches = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/pitches/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    selected_pitch_id = generated_pitches["snapshot"]["pitch_batches"][0]["pitches"][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/pitch",
            json={"pitch_id": selected_pitch_id, "origin": "workspace"},
        ).status_code
        == 200
    )
    generated_characters = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/characters/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    selected_character_sheet_id = generated_characters["snapshot"]["character_sheet_batches"][0][
        "character_sheets"
    ][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/character-sheet",
            json={
                "character_sheet_id": selected_character_sheet_id,
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    generated_beats = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/beats/generate",
        json={"origin": "workspace"},
    ).json()
    first_beat_sheet_id = generated_beats["snapshot"]["beat_sheet_revisions"][0]["id"]

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/selections/beat-sheet",
        json={
            "beat_sheet_id": first_beat_sheet_id,
            "origin": "workspace",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "selection.recorded"
    assert payload["event"]["payload"]["selection_kind"] == "beat_sheet"
    assert payload["snapshot"]["current_stage"] == "story_setup"
    assert payload["snapshot"]["resume_stage"] == "story_setup"
    assert payload["snapshot"]["selected_beat_sheet"]["id"] == first_beat_sheet_id
    assert payload["snapshot"]["selected_beat_sheet"]["is_selected"] is True
    assert payload["snapshot"]["stage_states"][5]["status"] == "completed"


def test_save_session_story_setup_endpoint_persists_soft_targets_and_returns_selection_event(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    created = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Story Setup API"},
    ).json()

    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/genre",
            json={"genre_slug": "quest-fantasy", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/tone",
            json={"tone_profile_slug": "hushed-wonder", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/story-brief",
            json={
                "story_idea": (
                    "A child follows floating lanterns through a harbor and helps a shy otter "
                    "guardian guide each light back to bed."
                ),
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    generated_pitches = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/pitches/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    selected_pitch_id = generated_pitches["snapshot"]["pitch_batches"][0]["pitches"][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/pitch",
            json={"pitch_id": selected_pitch_id, "origin": "workspace"},
        ).status_code
        == 200
    )
    generated_characters = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/characters/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    selected_character_sheet_id = generated_characters["snapshot"]["character_sheet_batches"][0][
        "character_sheets"
    ][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/character-sheet",
            json={
                "character_sheet_id": selected_character_sheet_id,
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    generated_beats = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/beats/generate",
        json={"origin": "workspace"},
    ).json()
    first_beat_sheet_id = generated_beats["snapshot"]["beat_sheet_revisions"][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/beat-sheet",
            json={
                "beat_sheet_id": first_beat_sheet_id,
                "origin": "workspace",
            },
        ).status_code
        == 200
    )

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/story-setup",
        json={
            "target_word_count": 1800,
            "target_runtime_minutes": 12,
            "chapter_count": 3,
            "approximate_scene_count": 8,
            "guidance_notes": "Treat these as calm planning targets, not promises.",
            "origin": "workspace",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "selection.recorded"
    assert payload["event"]["stage"] == "story_setup"
    assert payload["event"]["payload"]["selection_kind"] == "story_setup"
    assert payload["event"]["payload"]["source"] == "workspace"
    assert payload["snapshot"]["current_stage"] == "composition"
    assert payload["snapshot"]["resume_stage"] == "composition"
    assert payload["snapshot"]["selected_story_setup"]["target_word_count"] == 1800
    assert payload["snapshot"]["selected_story_setup"]["target_runtime_minutes"] == 12
    assert payload["snapshot"]["selected_story_setup"]["chapter_count"] == 3
    assert payload["snapshot"]["selected_story_setup"]["approximate_scene_count"] == 8
    assert payload["snapshot"]["selected_story_setup"]["guidance_notes"] == (
        "Treat these as calm planning targets, not promises."
    )
    assert payload["snapshot"]["selected_story_outline"]["outline_kind"] == "chapter"
    assert len(payload["snapshot"]["selected_story_outline"]["cards"]) == 3
    assert payload["snapshot"]["stage_states"][6]["status"] == "completed"


def test_save_session_story_outline_endpoint_persists_outline_revision(
    session_api_client: TestClient,
) -> None:
    _seed_catalog_rows()
    created = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Outline Edit"},
    ).json()
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/genre",
            json={"genre_slug": "quest-fantasy", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/tone",
            json={"tone_profile_slug": "hushed-wonder", "origin": "workspace"},
        ).status_code
        == 200
    )
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/story-brief",
            json={
                "raw_brief": (
                    "A harbor child and an otter guide the last lantern home before bed."
                ),
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    generated_pitches = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/pitches/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    selected_pitch_id = generated_pitches["snapshot"]["pitch_batches"][0]["pitches"][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/pitch",
            json={"pitch_id": selected_pitch_id, "origin": "workspace"},
        ).status_code
        == 200
    )
    generated_characters = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/characters/generate",
        json={"candidate_count": 3, "origin": "workspace"},
    ).json()
    selected_character_sheet_id = generated_characters["snapshot"]["character_sheet_batches"][0][
        "character_sheets"
    ][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/character-sheet",
            json={
                "character_sheet_id": selected_character_sheet_id,
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    generated_beats = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/beats/generate",
        json={"origin": "workspace"},
    ).json()
    selected_beat_sheet_id = generated_beats["snapshot"]["beat_sheet_revisions"][0]["id"]
    assert (
        session_api_client.post(
            f"/api/v1/sessions/{created['id']}/selections/beat-sheet",
            json={
                "beat_sheet_id": selected_beat_sheet_id,
                "origin": "workspace",
            },
        ).status_code
        == 200
    )
    story_setup_payload = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/story-setup",
        json={
            "target_word_count": 1800,
            "target_runtime_minutes": 12,
            "chapter_count": 3,
            "approximate_scene_count": 8,
            "guidance_notes": "Keep every chapter ending visibly softer than it began.",
            "origin": "workspace",
        },
    ).json()
    selected_outline = story_setup_payload["snapshot"]["selected_story_outline"]
    cards = selected_outline["cards"]
    cards[0]["summary"] = (
        "Open with a calmer harbor image, then move Mira toward the first drifting lantern."
    )

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/story-outline",
        json={
            "outline_id": selected_outline["id"],
            "summary": selected_outline["summary"],
            "cards": cards,
            "origin": "workspace",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "content.user_edit.recorded"
    assert payload["event"]["stage"] == "story_setup"
    assert payload["event"]["payload"]["target_kind"] == "story_outline"
    assert payload["snapshot"]["selected_story_outline"]["revision_number"] == 2
    assert payload["snapshot"]["selected_story_outline"]["cards"][0]["summary"].startswith(
        "Open with a calmer harbor image"
    )


def test_composition_control_endpoints_manage_durable_job_state(
    session_api_client: TestClient,
) -> None:
    seeded = _create_composition_ready_session_via_api(session_api_client)
    session_id = str(seeded["session_id"])

    start_response = session_api_client.post(
        f"/api/v1/sessions/{session_id}/composition/start",
        json={"mode": "fresh", "origin": "workspace"},
    )

    assert start_response.status_code == 200
    start_payload = start_response.json()
    composition_job_id = start_payload["job"]["id"]
    composition_stage = next(
        stage
        for stage in start_payload["snapshot"]["stage_states"]
        if stage["stage"] == "composition"
    )
    assert start_payload["job"]["status"] == "queued"
    assert start_payload["job"]["total_segments"] == 3
    assert start_payload["event"]["event_type"] == "composition.progress.recorded"
    assert composition_stage["status"] == "in_progress"

    pause_response = session_api_client.post(
        f"/api/v1/sessions/{session_id}/composition/{composition_job_id}/pause"
    )
    assert pause_response.status_code == 200
    pause_payload = pause_response.json()
    paused_stage = next(
        stage
        for stage in pause_payload["snapshot"]["stage_states"]
        if stage["stage"] == "composition"
    )
    assert pause_payload["job"]["status"] == "paused"
    assert paused_stage["detail"].startswith("Writing paused at")

    resume_response = session_api_client.post(
        f"/api/v1/sessions/{session_id}/composition/{composition_job_id}/resume"
    )
    assert resume_response.status_code == 200
    resume_payload = resume_response.json()
    resumed_stage = next(
        stage
        for stage in resume_payload["snapshot"]["stage_states"]
        if stage["stage"] == "composition"
    )
    assert resume_payload["job"]["status"] == "queued"
    assert resumed_stage["detail"].startswith("Queued writing to resume at segment")

    cancel_response = session_api_client.post(
        f"/api/v1/sessions/{session_id}/composition/{composition_job_id}/cancel"
    )
    assert cancel_response.status_code == 200
    cancel_payload = cancel_response.json()
    cancelled_stage = next(
        stage
        for stage in cancel_payload["snapshot"]["stage_states"]
        if stage["stage"] == "composition"
    )
    assert cancel_payload["job"]["status"] == "cancelled"
    assert cancelled_stage["status"] == "needs_regeneration"


def test_redirect_composition_endpoint_queues_rewrite_job(
    session_api_client: TestClient,
) -> None:
    seeded = _create_composition_ready_session_via_api(session_api_client)
    session_id = str(seeded["session_id"])

    start_payload = session_api_client.post(
        f"/api/v1/sessions/{session_id}/composition/start",
        json={"mode": "fresh", "origin": "workspace"},
    ).json()
    original_job_id = start_payload["job"]["id"]

    redirect_response = session_api_client.post(
        f"/api/v1/sessions/{session_id}/composition/{original_job_id}/redirect",
        json={
            "instructions": "Soften the midpoint and make the helper visible earlier.",
            "rewrite_from_segment_index": 2,
            "origin": "workspace",
        },
    )

    assert redirect_response.status_code == 200
    payload = redirect_response.json()
    assert payload["job"]["id"] != original_job_id
    assert payload["job"]["job_kind"] == "rewrite"
    assert payload["job"]["current_segment_index"] == 2
    assert payload["event"]["payload"]["job_id"] == payload["job"]["id"]


def test_reject_rewrite_endpoint_keeps_the_current_segment_version(
    session_api_client: TestClient,
) -> None:
    seeded = _create_composition_ready_session_via_api(session_api_client)
    session_id = str(seeded["session_id"])

    with get_session_factory()() as db_session:
        draft_job = CompositionJob(
            session_id=session_id,
            job_kind=CompositionJobKind.DRAFT,
            status=JobStatus.COMPLETED,
            progress_percent=100,
            current_segment_index=2,
            metadata_json={
                "accepted_story_so_far": (
                    "Draft segment 1 settles the harbor.\n\n"
                    "Draft segment 2 keeps the bell close."
                ),
                "latest_partial_output": (
                    "Draft segment 1 settles the harbor.\n\n"
                    "Draft segment 2 keeps the bell close."
                ),
                "total_segments": 2,
                "start_segment_index": 1,
            },
        )
        db_session.add(draft_job)
        db_session.flush()

        draft_segment_one = CompositionSegment(
            session_id=session_id,
            composition_job_id=draft_job.id,
            segment_index=1,
            revision_number=1,
            status=JobStatus.COMPLETED,
            acceptance_state=CompositionSegmentAcceptanceState.ACCEPTED,
            accepted_text="Draft segment 1 settles the harbor.",
            text_content="Draft segment 1 settles the harbor.",
            word_count=6,
            payload={"outline_card_title": "Opening harbor"},
        )
        draft_segment_two = CompositionSegment(
            session_id=session_id,
            composition_job_id=draft_job.id,
            segment_index=2,
            revision_number=1,
            status=JobStatus.COMPLETED,
            acceptance_state=CompositionSegmentAcceptanceState.ACCEPTED,
            accepted_text="Draft segment 2 keeps the bell close.",
            text_content="Draft segment 2 keeps the bell close.",
            word_count=7,
            payload={"outline_card_title": "Lantern cove"},
        )
        db_session.add_all([draft_segment_one, draft_segment_two])
        db_session.flush()

        rewrite_job = CompositionJob(
            session_id=session_id,
            job_kind=CompositionJobKind.REWRITE,
            status=JobStatus.COMPLETED,
            progress_percent=100,
            current_segment_index=2,
            rewrite_to_segment_index=2,
            metadata_json={
                "accepted_story_so_far": (
                    "Draft segment 1 settles the harbor.\n\n"
                    "Draft segment 2 keeps the bell close."
                ),
                "latest_partial_output": (
                    "Draft segment 1 settles the harbor.\n\n"
                    "Draft segment 2 keeps the bell close."
                ),
                "total_segments": 1,
                "start_segment_index": 2,
            },
        )
        db_session.add(rewrite_job)
        db_session.flush()

        rewrite_segment = CompositionSegment(
            session_id=session_id,
            composition_job_id=rewrite_job.id,
            segment_index=2,
            revision_number=2,
            status=JobStatus.COMPLETED,
            acceptance_state=CompositionSegmentAcceptanceState.PENDING,
            accepted_text="Rewrite segment 2 opens the cove more gently.",
            text_content="Rewrite segment 2 opens the cove more gently.",
            word_count=8,
            payload={"outline_card_title": "Lantern cove"},
        )
        db_session.add(rewrite_segment)
        db_session.commit()

    response = session_api_client.post(
        f"/api/v1/sessions/{session_id}/composition/{rewrite_job.id}/reject",
        json={"origin": "workspace"},
    )

    assert response.status_code == 200
    payload = response.json()
    composition_stage = next(
        stage
        for stage in payload["snapshot"]["stage_states"]
        if stage["stage"] == "composition"
    )
    segment_two = next(
        segment
        for segment in payload["snapshot"]["composition_segments"]
        if segment["segment_index"] == 2
    )

    assert payload["job"]["id"] == rewrite_job.id
    assert payload["job"]["pending_review"] is False
    assert payload["event"]["event_type"] == "content.user_edit.recorded"
    assert (
        payload["event"]["payload"]["summary_text"]
        == "Kept the current manuscript and dismissed the rewrite candidate."
    )
    assert composition_stage["status"] == "completed"
    assert segment_two["current_revision_number"] == 1
    assert segment_two["pending_version_id"] is None
    assert [
        version["acceptance_state"] for version in segment_two["versions"]
    ] == ["rejected", "accepted"]


def test_session_events_websocket_replays_composition_job_status(
    session_api_client: TestClient,
) -> None:
    seeded = _create_composition_ready_session_via_api(session_api_client)
    session_id = str(seeded["session_id"])
    latest_sequence_number = session_api_client.get(
        f"/api/v1/sessions/{session_id}/history?limit=1"
    ).json()["latest_sequence_number"]

    session_api_client.post(
        f"/api/v1/sessions/{session_id}/composition/start",
        json={"mode": "fresh", "origin": "workspace"},
    )

    with session_api_client.websocket_connect("/api/v1/sessions/events/ws") as websocket:
        websocket.send_json(
            {
                "schema_version": 1,
                "action": "subscribe",
                "session_id": session_id,
                "tab_id": "tab-1",
                "last_sequence_number": latest_sequence_number,
            }
        )
        ack = websocket.receive_json()

        replayed_job_status = None
        for _ in range(3):
            message = websocket.receive_json()
            if message["type"] == "job.status" and message["payload"]["job_kind"] == "composition":
                replayed_job_status = message
                break

    assert ack["action"] == "subscribed"
    assert ack["channel"] == f"session:{session_id}"
    assert replayed_job_status is not None
    assert replayed_job_status["delivery"] == "replay"
    assert replayed_job_status["payload"]["status"] == "queued"
    assert replayed_job_status["payload"]["total_segments"] == 3


def test_session_realtime_service_streams_composition_chunk_deltas(
    session_api_client: TestClient,
) -> None:
    seeded = _create_composition_ready_session_via_api(session_api_client)
    session_id = str(seeded["session_id"])
    start_payload = session_api_client.post(
        f"/api/v1/sessions/{session_id}/composition/start",
        json={"mode": "fresh", "origin": "workspace"},
    ).json()
    composition_job_id = start_payload["job"]["id"]

    with get_session_factory()() as db_session:
        job = db_session.get(CompositionJob, composition_job_id)
        segment = db_session.execute(
            select(CompositionSegment)
            .where(CompositionSegment.composition_job_id == composition_job_id)
            .order_by(CompositionSegment.segment_index.asc())
            .limit(1)
        ).scalar_one()

        assert job is not None
        job.status = JobStatus.IN_PROGRESS
        job.progress_percent = 12
        job.metadata_json = {
            **(job.metadata_json if isinstance(job.metadata_json, dict) else {}),
            "accepted_story_so_far": "",
            "latest_partial_output": "",
            "current_segment_id": segment.id,
        }
        segment.status = JobStatus.IN_PROGRESS
        db_session.commit()

        realtime = SessionRealtimeService(db_session)
        initial_events, cursor = realtime.read_composition_chunk_events(
            session_id,
            cursor=CompositionChunkCursor(),
        )

        streamed_text = (
            "Mira followed the bell toward the quieter cove while Pip kept the "
            "water calm and readable."
        )
        job.metadata_json = {
            **(job.metadata_json if isinstance(job.metadata_json, dict) else {}),
            "accepted_story_so_far": streamed_text,
            "latest_partial_output": streamed_text,
            "latest_segment_summary": (
                "Segment 1 moved Mira into the quieter cove without losing the calm tone."
            ),
            "current_segment_id": segment.id,
        }
        segment.accepted_text = streamed_text
        segment.text_content = streamed_text
        db_session.commit()

        streamed_events, next_cursor = realtime.read_composition_chunk_events(
            session_id,
            cursor=cursor,
        )

    delta_events = [
        event
        for event in streamed_events
        if event.payload.chunk_kind.value == "delta"
    ]
    summary_events = [
        event
        for event in streamed_events
        if event.payload.chunk_kind.value == "segment_summary"
    ]

    assert len(initial_events) == 1
    assert initial_events[0].payload.chunk_kind.value == "segment_start"
    assert delta_events
    assert "".join(event.payload.text or "" for event in delta_events) == streamed_text
    assert summary_events[0].payload.summary.startswith("Segment 1 moved Mira")
    assert next_cursor.story_text == streamed_text


def test_session_realtime_service_suppresses_initial_chunk_baseline_and_keeps_snapshot_recoverable(
    session_api_client: TestClient,
) -> None:
    seeded = _create_composition_ready_session_via_api(session_api_client)
    session_id = str(seeded["session_id"])
    start_payload = session_api_client.post(
        f"/api/v1/sessions/{session_id}/composition/start",
        json={"mode": "fresh", "origin": "workspace"},
    ).json()
    composition_job_id = start_payload["job"]["id"]
    accepted_story = (
        "Draft segment 1 settles the harbor.\n\nMira followed the bell toward the quieter cove."
    )

    with get_session_factory()() as db_session:
        job = db_session.get(CompositionJob, composition_job_id)
        segment = db_session.execute(
            select(CompositionSegment)
            .where(CompositionSegment.composition_job_id == composition_job_id)
            .order_by(CompositionSegment.segment_index.asc())
            .limit(1)
        ).scalar_one()

        assert job is not None
        job.status = JobStatus.IN_PROGRESS
        job.progress_percent = 38
        job.metadata_json = {
            **(job.metadata_json if isinstance(job.metadata_json, dict) else {}),
            "accepted_story_so_far": accepted_story,
            "latest_partial_output": accepted_story,
            "latest_segment_summary": "Segment 1 settled the harbor before the cove opened.",
            "current_segment_id": segment.id,
        }
        segment.status = JobStatus.IN_PROGRESS
        segment.accepted_text = accepted_story
        segment.text_content = accepted_story
        db_session.commit()

        snapshot = SessionService(db_session).load_session_snapshot(session_id)
        realtime = SessionRealtimeService(db_session)
        initial_events, cursor = realtime.read_composition_chunk_events(
            session_id,
            cursor=CompositionChunkCursor(),
        )

    assert snapshot.active_composition_job is not None
    assert snapshot.active_composition_job.accepted_story_so_far == accepted_story
    assert len(initial_events) == 1
    assert initial_events[0].payload.chunk_kind.value == "segment_start"
    assert cursor.story_text == accepted_story


def test_hydrate_session_endpoint_preserves_chat_navigation_bridge_history(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Bridge Replay"},
    )
    created = create_response.json()

    parse_response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/chat/intents",
        json={
            "message": "/next-stage",
            "explicit_command": {
                "command_id": "next_stage",
                "source": "quick_action",
                "proposed_actions": {
                    "schema_version": 1,
                    "actions": [
                        {
                            "schema_version": 1,
                            "action_type": "navigate_to_stage",
                            "target_stage": "tone",
                            "confidence": 1,
                            "rationale": "Explicit command requested navigation to Tone.",
                            "requires_confirmation": False,
                            "extracted_values": {},
                        }
                    ],
                },
            },
        },
    )

    assert parse_response.status_code == 200
    assert parse_response.json()["policy_evaluation"]["evaluated_actions"][0]["decision"] == (
        "accepted"
    )

    ui_action_response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/ui-actions",
        json={
            "action": "navigate_to_stage",
            "stage": "tone",
            "control_id": "chat-intent",
            "value_summary": "Tone",
            "origin": "chat",
        },
    )

    assert ui_action_response.status_code == 201

    hydration_response = session_api_client.get(f"/api/v1/sessions/{created['id']}/hydrate")

    assert hydration_response.status_code == 200
    payload = hydration_response.json()

    assert [event["event_type"] for event in payload["recent_history"]["events"]] == [
        "session.created",
        "chat.message.recorded",
        "chat.intent.parsed",
        "chat.message.recorded",
        "ui.action.recorded",
    ]
    assert payload["recent_history"]["events"][-1]["payload"] == {
        "schema_version": 1,
        "action": "navigate_to_stage",
        "control_id": "chat-intent",
        "value_summary": "Tone",
        "origin": "chat",
    }
    assert payload["snapshot"]["current_stage"] == "genre"
    assert payload["snapshot"]["resume_stage"] == "genre"
    assert payload["hydration"]["strategy"] == "materialized_only"
    assert payload["hydration"]["latest_sequence_number"] == 5
    assert payload["hydration"]["history_event_count"] == 5


def test_apply_session_context_update_endpoint_returns_updated_snapshot_and_event(
    session_api_client: TestClient,
) -> None:
    db_session = get_session_factory()()
    try:
        snapshot = SessionService(db_session).create_session(working_title="Context Update")
        service = SessionService(db_session)
        for stage in (
            WorkflowStage.GENRE,
            WorkflowStage.TONE,
            WorkflowStage.BRIEF,
            WorkflowStage.PITCHES,
            WorkflowStage.CHARACTERS,
            WorkflowStage.BEATS,
        ):
            service.update_stage_state(
                snapshot.id,
                stage=stage,
                status=WorkflowStageState.COMPLETED,
                detail=f"Accepted {stage.value}.",
            )
    finally:
        db_session.close()

    response = session_api_client.post(
        f"/api/v1/sessions/{snapshot.id}/context-updates",
        json={
            "target_kind": "stage_note",
            "stage": "beats",
            "control_id": "stage-note-editor",
            "origin": "workspace",
            "values": {
                "detail": "Add one calmer beat before the return home.",
            },
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "content.user_edit.recorded"
    assert payload["event"]["payload"]["field_values"] == {
        "detail": "Add one calmer beat before the return home.",
        "control_id": "stage-note-editor",
    }
    assert payload["snapshot"]["stage_states"][5]["detail"] == (
        "Add one calmer beat before the return home."
    )
    assert payload["snapshot"]["stage_states"][7]["status"] == "draft"
    assert (
        "Latest saved UI detail: Beat sheet: Add one calmer beat"
        in payload["snapshot"]["agent_context_summary"]
    )


def test_hydrate_session_endpoint_replays_context_updates_into_resumed_snapshot(
    session_api_client: TestClient,
) -> None:
    db_session = get_session_factory()()
    try:
        service = SessionService(db_session)
        snapshot = service.create_session(working_title="Replay Context")
        for stage in (
            WorkflowStage.GENRE,
            WorkflowStage.TONE,
            WorkflowStage.BRIEF,
            WorkflowStage.PITCHES,
            WorkflowStage.CHARACTERS,
            WorkflowStage.BEATS,
            WorkflowStage.STORY_SETUP,
            WorkflowStage.COMPOSITION,
            WorkflowStage.AUDIO,
        ):
            service.update_stage_state(
                snapshot.id,
                stage=stage,
                status=WorkflowStageState.COMPLETED,
                detail=f"Accepted {stage.value}.",
            )
    finally:
        db_session.close()

    ui_action_response = session_api_client.post(
        f"/api/v1/sessions/{snapshot.id}/ui-actions",
        json={
            "action": "navigate_to_stage",
            "stage": "audio",
            "control_id": "stage-navigator",
            "value_summary": "Audio",
            "origin": "workspace",
        },
    )
    assert ui_action_response.status_code == 201

    context_update_response = session_api_client.post(
        f"/api/v1/sessions/{snapshot.id}/context-updates",
        json={
            "target_kind": "stage_note",
            "stage": "beats",
            "control_id": "stage-note-editor",
            "origin": "workspace",
            "values": {
                "detail": "Soften the midpoint and let the return home land faster.",
            },
        },
    )
    assert context_update_response.status_code == 200

    hydration_response = session_api_client.get(f"/api/v1/sessions/{snapshot.id}/hydrate")

    assert hydration_response.status_code == 200
    payload = hydration_response.json()

    assert [event["event_type"] for event in payload["recent_history"]["events"][-3:]] == [
        "ui.action.recorded",
        "workflow.stage_changed",
        "content.user_edit.recorded",
    ]
    assert payload["snapshot"]["stage_states"][5]["detail"] == (
        "Soften the midpoint and let the return home land faster."
    )
    assert payload["snapshot"]["stage_states"][6]["status"] == "needs_regeneration"
    assert payload["snapshot"]["stage_states"][7]["status"] == "needs_regeneration"
    assert payload["snapshot"]["stage_states"][8]["status"] == "needs_regeneration"
    assert payload["snapshot"]["resume_stage"] == "story_setup"
    assert payload["snapshot"]["overall_status"] == "needs_regeneration"
    assert payload["hydration"]["history_event_count"] == 13
    assert payload["hydration"]["latest_sequence_number"] == 13


def test_save_audio_settings_endpoint_persists_durable_audio_preferences(
    session_api_client: TestClient,
) -> None:
    db_session = get_session_factory()()
    now = datetime.now(timezone.utc)
    try:
        service = SessionService(db_session)
        snapshot = service.create_session(working_title="Audio API")
        for stage in (
            WorkflowStage.GENRE,
            WorkflowStage.TONE,
            WorkflowStage.BRIEF,
            WorkflowStage.PITCHES,
            WorkflowStage.CHARACTERS,
            WorkflowStage.BEATS,
            WorkflowStage.STORY_SETUP,
            WorkflowStage.COMPOSITION,
            WorkflowStage.AUDIO,
            WorkflowStage.FINALIZE,
        ):
            service.update_stage_state(
                snapshot.id,
                stage=stage,
                status=WorkflowStageState.COMPLETED,
                detail=f"Accepted {stage.value}.",
            )

        db_session.add(
            StorySetup(
                session_id=snapshot.id,
                revision_number=1,
                target_word_count=1800,
                target_runtime_minutes=12,
                is_selected=True,
                accepted_at=now,
            )
        )
        db_session.flush()

        audio_job = AudioJob(
            session_id=snapshot.id,
            status=JobStatus.COMPLETED,
            voice_key="moonbeam",
            playback_speed=0.95,
            include_background_music=False,
            music_profile="lullaby_piano",
            completed_at=now,
        )
        db_session.add(audio_job)
        db_session.flush()

        db_session.add(
            SessionAsset(
                session_id=snapshot.id,
                audio_job_id=audio_job.id,
                asset_kind=AssetKind.FINAL_AUDIO,
                status=AssetStatus.READY,
                storage_bucket="storyteller-exports",
                object_path="sessions/audio-api/story.mp3",
                mime_type="audio/mpeg",
                byte_size=8192,
                ready_at=now,
            )
        )
        db_session.commit()
    finally:
        db_session.close()

    response = session_api_client.post(
        f"/api/v1/sessions/{snapshot.id}/audio-settings",
        json={
            "voice_key": "hearthside",
            "narration_style": "warm",
            "playback_speed": 0.85,
            "include_background_music": True,
            "music_profile": "night_ambience",
            "narration_volume": 88,
            "music_volume": 18,
            "guidance_notes": "Ease off even more during the final chapter.",
            "origin": "workspace",
        },
    )

    assert response.status_code == 200
    payload = response.json()

    assert payload["event"]["event_type"] == "content.user_edit.recorded"
    assert payload["event"]["stage"] == "audio"
    assert payload["event"]["payload"]["changed_fields"] == [
        "voice_key",
        "narration_style",
        "playback_speed",
        "include_background_music",
        "music_profile",
        "narration_volume",
        "music_volume",
        "guidance_notes",
    ]
    assert payload["event"]["payload"]["field_values"] == {
        "voice_key": "hearthside",
        "narration_style": "warm",
        "playback_speed": 0.85,
        "include_background_music": True,
        "music_profile": "night_ambience",
        "narration_volume": 88,
        "music_volume": 18,
        "guidance_notes": "Ease off even more during the final chapter.",
    }
    assert payload["snapshot"]["audio_settings"]["voice_key"] == "hearthside"
    assert payload["snapshot"]["audio_settings"]["narration_style"] == "warm"
    assert payload["snapshot"]["audio_settings"]["playback_speed"] == 0.85
    assert payload["snapshot"]["audio_settings"]["include_background_music"] is True
    assert payload["snapshot"]["audio_settings"]["music_profile"] == "night_ambience"
    assert payload["snapshot"]["audio_settings"]["narration_volume"] == 88
    assert payload["snapshot"]["audio_settings"]["music_volume"] == 18
    assert len(payload["snapshot"]["audio_settings"]["music_profile_options"]) == 3
    assert (
        payload["snapshot"]["audio_settings"]["mix_preview"]["strategy"]
        == "curated_bed_ducked"
    )
    assert (
        payload["snapshot"]["audio_settings"]["runtime_estimate"]["basis_source"]
        == "story_setup_target"
    )
    stage_map = {
        stage["stage"]: stage
        for stage in payload["snapshot"]["stage_states"]
    }
    assert stage_map["audio"]["status"] == "needs_regeneration"
    assert stage_map["finalize"]["status"] == "needs_regeneration"


def test_get_session_snapshot_endpoint_returns_404_for_missing_session(
    session_api_client: TestClient,
) -> None:
    response = session_api_client.get("/api/v1/sessions/missing-session")

    assert response.status_code == 404
    assert response.json() == {
        "detail": "session 'missing-session' was not found",
    }


def test_hydrate_session_endpoint_returns_404_for_missing_session(
    session_api_client: TestClient,
) -> None:
    response = session_api_client.get("/api/v1/sessions/missing-session/hydrate")

    assert response.status_code == 404
    assert response.json() == {
        "detail": "session 'missing-session' was not found",
    }


def test_record_session_ui_action_endpoint_returns_404_for_missing_session(
    session_api_client: TestClient,
) -> None:
    response = session_api_client.post(
        "/api/v1/sessions/missing-session/ui-actions",
        json={
            "action": "navigate_to_stage",
            "stage": "audio",
        },
    )

    assert response.status_code == 404
    assert response.json() == {
        "detail": "session 'missing-session' was not found",
    }


def test_apply_session_context_update_endpoint_returns_422_for_unsupported_stage(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Unsupported Stage"},
    )
    created = create_response.json()

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/context-updates",
        json={
            "target_kind": "stage_note",
            "stage": "genre",
            "values": {
                "detail": "Quest fantasy should lean quieter.",
            },
        },
    )

    assert response.status_code == 422
    assert "does not support durable note edits" in response.json()["detail"]


def test_save_session_story_setup_endpoint_returns_422_before_beat_selection(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Setup Guard"},
    )
    created = create_response.json()

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/story-setup",
        json={
            "target_runtime_minutes": 9,
            "origin": "workspace",
        },
    )

    assert response.status_code == 422
    assert "requires a selected beat sheet" in response.json()["detail"]


def test_apply_session_context_update_endpoint_returns_409_for_invalid_transition(
    session_api_client: TestClient,
) -> None:
    create_response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "Transition Guard"},
    )
    created = create_response.json()

    response = session_api_client.post(
        f"/api/v1/sessions/{created['id']}/context-updates",
        json={
            "target_kind": "stage_note",
            "stage": "story_setup",
            "values": {
                "detail": "Target a shorter read-aloud.",
            },
        },
    )

    assert response.status_code == 409
    assert (
        "cannot set 'story_setup' to 'in_progress' before prerequisites are completed"
        in (response.json()["detail"])
    )


def test_create_session_endpoint_returns_a_fresh_snapshot(
    session_api_client: TestClient,
) -> None:
    response = session_api_client.post(
        "/api/v1/sessions",
        json={"working_title": "  Moonlit Harbor  "},
    )

    assert response.status_code == 201
    payload = response.json()

    assert payload["display_title"] == "Moonlit Harbor"
    assert payload["working_title"] == "Moonlit Harbor"
    assert payload["resume_stage"] == "genre"
    assert payload["overall_status"] == "draft"
    assert payload["progress"] == {
        "total_stages": 10,
        "completed_stages": 0,
        "in_progress_stages": 0,
        "needs_regeneration_stages": 0,
    }
