from __future__ import annotations

import wave
from io import BytesIO
from typing import Iterator

import pytest
from app.db.session import get_engine, get_session_factory
from app.main import create_app
from app.services.catalog import CATALOG_FILE_PATH, load_catalog_document, seed_catalog
from app.settings import AppSettings, get_settings
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session, sessionmaker
from tests.support.audio_pipeline import SequenceTextToSpeechAdapter, build_synthesis_result
from tests.support.full_journey import (
    FullJourneyBeatSheetGenerationAdapter,
    FullJourneyBriefNormalizationAdapter,
    FullJourneyCharacterGenerationAdapter,
    FullJourneyCompositionWriter,
    FullJourneyPitchGenerationAdapter,
    build_worker,
    run_full_journey_planning,
    run_worker_until_idle,
)

pytestmark = pytest.mark.integration


@pytest.fixture
def full_journey_api_client(
    monkeypatch: pytest.MonkeyPatch,
    integration_settings: AppSettings,
    db_session_factory: sessionmaker[Session],
) -> Iterator[TestClient]:
    _apply_integration_settings(monkeypatch, integration_settings)
    get_settings.cache_clear()
    get_engine.cache_clear()
    get_session_factory.cache_clear()

    with db_session_factory() as session:
        seed_catalog(session, load_catalog_document(CATALOG_FILE_PATH))

    app = create_app()
    app.state.brief_normalization_adapter = FullJourneyBriefNormalizationAdapter()
    app.state.pitch_generation_adapter = FullJourneyPitchGenerationAdapter()
    app.state.character_generation_adapter = FullJourneyCharacterGenerationAdapter()
    app.state.beat_sheet_generation_adapter = FullJourneyBeatSheetGenerationAdapter()

    try:
        with TestClient(app) as client:
            yield client
    finally:
        get_settings.cache_clear()
        get_engine.cache_clear()
        get_session_factory.cache_clear()


def test_full_journey_e2e_persists_story_audio_exports_and_resume_state(
    full_journey_api_client: TestClient,
    db_session_factory: sessionmaker[Session],
    object_storage,
) -> None:
    client = full_journey_api_client

    recent_before = _expect_json(client.get("/api/v1/sessions"))
    assert recent_before == []

    journey = run_full_journey_planning(client)
    session_id = journey.session_id

    recent_during_planning = _expect_json(client.get("/api/v1/sessions"))
    assert recent_during_planning[0]["id"] == session_id
    assert recent_during_planning[0]["current_stage"] == "composition"
    assert recent_during_planning[0]["library_summary"]["artifact_readiness"] == {
        "story_text": "missing",
        "story_docx": "missing",
        "final_audio": "missing",
        "ready_count": 0,
        "total_count": 3,
    }

    refreshed_planning = _expect_json(client.get(f"/api/v1/sessions/{session_id}/hydrate"))
    assert refreshed_planning["snapshot"]["selected_pitch"]["generation_kind"] == "refinement"
    assert refreshed_planning["snapshot"]["selected_character_sheet"]["generation_kind"] == (
        "refinement"
    )
    assert refreshed_planning["snapshot"]["selected_story_setup"]["target_runtime_minutes"] == 15
    assert refreshed_planning["snapshot"]["selected_story_outline"]["chapter_count"] == 4

    composition_started = _expect_json(
        client.post(
            f"/api/v1/sessions/{session_id}/composition/start",
            json={"mode": "fresh", "origin": "workspace"},
        )
    )
    assert composition_started["job"]["status"] == "queued"
    assert (
        composition_started["snapshot"]["active_composition_job"]["id"]
        == (composition_started["job"]["id"])
    )

    composition_worker = build_worker(
        session_factory=db_session_factory,
        object_storage=object_storage,
        composition_writer=FullJourneyCompositionWriter(),
        tts_adapter=SequenceTextToSpeechAdapter(
            [build_synthesis_result(sample_value=index) for index in range(1, 5)]
        ),
    )
    composition_worker_jobs = run_worker_until_idle(composition_worker)
    assert composition_worker_jobs >= 4

    hydrated_after_composition = _expect_json(client.get(f"/api/v1/sessions/{session_id}/hydrate"))
    snapshot_after_composition = hydrated_after_composition["snapshot"]
    assert snapshot_after_composition["active_composition_job"] is None
    assert snapshot_after_composition["latest_composition_job"]["status"] == "completed"
    assert snapshot_after_composition["latest_story_asset"]["asset_kind"] == "story_text"
    assert "Segment 4" in _download_text(
        client.get(f"/api/v1/sessions/{session_id}/artifacts/story-text")
    )

    inventory_after_composition = _expect_json(
        client.get(f"/api/v1/sessions/{session_id}/artifacts")
    )
    assert _artifact_item(inventory_after_composition, "story_text")["status"] == "ready"
    assert _artifact_item(inventory_after_composition, "story_docx")["status"] == "missing"
    assert _artifact_item(inventory_after_composition, "final_audio")["status"] == "missing"

    docx_asset = _expect_json(client.post(f"/api/v1/sessions/{session_id}/artifacts/story-docx"))
    assert docx_asset["asset_kind"] == "story_docx"

    audio_settings = _expect_json(
        client.post(
            f"/api/v1/sessions/{session_id}/audio-settings",
            json={
                "voice_key": "hearthside",
                "narration_style": "warm",
                "playback_speed": 0.9,
                "include_background_music": False,
                "guidance_notes": "Keep the delivery warm, steady, and bedtime-soft.",
                "origin": "workspace",
            },
        )
    )
    assert audio_settings["snapshot"]["audio_settings"]["voice_key"] == "hearthside"

    audio_started = _expect_json(
        client.post(
            f"/api/v1/sessions/{session_id}/audio/generate",
            json={},
        )
    )
    assert audio_started["snapshot"]["active_audio_job"]["status"] == "queued"
    assert audio_started["event"]["stage"] == "audio"

    audio_worker = build_worker(
        session_factory=db_session_factory,
        object_storage=object_storage,
        composition_writer=FullJourneyCompositionWriter(),
        tts_adapter=SequenceTextToSpeechAdapter(
            [build_synthesis_result(sample_value=index) for index in range(1, 5)]
        ),
    )
    audio_worker_jobs = run_worker_until_idle(audio_worker)
    assert audio_worker_jobs >= 1

    final_snapshot = _expect_json(client.get(f"/api/v1/sessions/{session_id}"))
    final_hydration = _expect_json(client.get(f"/api/v1/sessions/{session_id}/hydrate"))
    final_inventory = _expect_json(client.get(f"/api/v1/sessions/{session_id}/artifacts"))
    history = _expect_json(client.get(f"/api/v1/sessions/{session_id}/history?limit=200"))

    assert final_snapshot["latest_story_asset"]["asset_kind"] == "story_text"
    assert final_snapshot["latest_story_export_asset"]["asset_kind"] == "story_docx"
    assert final_snapshot["latest_audio_asset"]["asset_kind"] == "final_audio"
    assert final_snapshot["active_audio_job"] is None
    assert final_snapshot["latest_audio_job"]["status"] == "completed"
    assert (
        final_hydration["snapshot"]["latest_audio_asset"]["id"]
        == (final_snapshot["latest_audio_asset"]["id"])
    )
    assert _artifact_item(final_inventory, "story_text")["status"] == "ready"
    assert _artifact_item(final_inventory, "story_docx")["status"] == "ready"
    assert _artifact_item(final_inventory, "final_audio")["status"] == "ready"

    docx_response = client.get(f"/api/v1/sessions/{session_id}/artifacts/story-docx")
    docx_response.raise_for_status()
    exported_docx = Document(BytesIO(docx_response.content))
    assert exported_docx.paragraphs[0].text == "Full Journey E2E"
    assert any(
        "Segment 1 lets Mira follow the silver bell" in paragraph.text
        for paragraph in exported_docx.paragraphs
    )

    audio_response = client.get(f"/api/v1/sessions/{session_id}/artifacts/final-audio")
    audio_response.raise_for_status()
    with wave.open(BytesIO(audio_response.content), "rb") as rendered_audio:
        assert rendered_audio.getnframes() > 0
        assert rendered_audio.getframerate() == 24_000

    recent_after_finalize = _expect_json(client.get("/api/v1/sessions"))
    assert recent_after_finalize[0]["id"] == session_id
    assert recent_after_finalize[0]["library_summary"]["artifact_readiness"] == {
        "story_text": "ready",
        "story_docx": "ready",
        "final_audio": "ready",
        "ready_count": 3,
        "total_count": 3,
    }
    assert recent_after_finalize[0]["library_summary"]["runtime_source"] == "final_audio"

    history_types = [event["event_type"] for event in history["events"]]
    assert "selection.recorded" in history_types
    assert "composition.progress.recorded" in history_types
    assert "audio.progress.recorded" in history_types


def _apply_integration_settings(
    monkeypatch: pytest.MonkeyPatch,
    settings: AppSettings,
) -> None:
    monkeypatch.setenv("STORYTELLER_SECRETS_FILE", "")
    monkeypatch.setenv("STORYTELLER_DATABASE_URL", settings.database.url)
    monkeypatch.setenv(
        "STORYTELLER_GEMINI_API_KEY",
        settings.gemini.api_key.get_secret_value(),
    )
    monkeypatch.setenv("STORYTELLER_GCS_ENDPOINT", settings.gcs.endpoint)
    monkeypatch.setenv("STORYTELLER_GCS_PROJECT_ID", settings.gcs.project_id)
    monkeypatch.setenv("STORYTELLER_GCS_PUBLIC_URL", settings.gcs.public_url or "")
    monkeypatch.setenv(
        "STORYTELLER_GCS_SESSIONS_BUCKET_NAME",
        settings.gcs.buckets.sessions,
    )
    monkeypatch.setenv(
        "STORYTELLER_GCS_AUDIO_BUCKET_NAME",
        settings.gcs.buckets.audio,
    )
    monkeypatch.setenv(
        "STORYTELLER_GCS_EXPORTS_BUCKET_NAME",
        settings.gcs.buckets.exports,
    )


def _artifact_item(
    inventory: dict[str, object],
    key: str,
) -> dict[str, object]:
    for item in inventory["items"]:
        if item["key"] == key:
            return item
    raise AssertionError(f"artifact {key!r} was not present in inventory")


def _download_text(response) -> str:
    response.raise_for_status()
    return response.content.decode("utf-8")


def _expect_json(response, expected_status: int = 200):
    assert response.status_code == expected_status, response.text
    return response.json()
