from __future__ import annotations

from io import BytesIO

from app.db import AssetKind, AssetStatus, Base, SessionAsset, StorySession, make_engine
from app.services.story_exports import StoryDocxExportService
from app.settings import get_settings
from docx import Document
from sqlalchemy.orm import sessionmaker
from tests.support.in_memory_storage import InMemoryObjectStorage


def test_story_docx_export_service_generates_and_reuses_matching_export() -> None:
    engine = make_engine("sqlite+pysqlite:///:memory:")
    Base.metadata.create_all(engine)
    db_session = sessionmaker(bind=engine, autoflush=False, expire_on_commit=False)()
    storage = InMemoryObjectStorage(get_settings())

    try:
        story_session = StorySession(working_title="Lantern Harbor")
        db_session.add(story_session)
        db_session.flush()

        story_location = storage.paths.export_asset(
            session_id=story_session.id,
            export_kind="story",
            export_id="accepted-manuscript",
            extension="md",
        )
        storage.upload_text(
            story_location,
            "# Chapter 1\n\nMira carried the lantern home.",
            content_type="text/markdown; charset=utf-8",
        )
        story_asset = SessionAsset(
            session_id=story_session.id,
            asset_kind=AssetKind.STORY_TEXT,
            status=AssetStatus.READY,
            storage_bucket=story_location.bucket,
            object_path=story_location.key,
            mime_type="text/markdown",
        )
        db_session.add(story_asset)
        db_session.commit()

        service = StoryDocxExportService(db_session, object_storage=storage)
        first_export = service.ensure_docx_asset(story_session.id)
        second_export = service.ensure_docx_asset(story_session.id)

        assert first_export.id == second_export.id
        assert first_export.asset_kind == AssetKind.STORY_DOCX
        assert first_export.metadata_json["source_story_asset_id"] == story_asset.id

        exported_bytes = storage.download_bytes(
            storage.paths.export_asset(
                session_id=story_session.id,
                export_kind="docx",
                export_id="final-manuscript",
                extension="docx",
            )
        )
        document = Document(BytesIO(exported_bytes))
        assert document.paragraphs[0].text == "Lantern Harbor"
        assert any(
            paragraph.text == "Mira carried the lantern home."
            for paragraph in document.paragraphs
        )
    finally:
        db_session.close()
        engine.dispose()


def test_story_docx_export_service_regenerates_when_story_text_changes() -> None:
    engine = make_engine("sqlite+pysqlite:///:memory:")
    Base.metadata.create_all(engine)
    db_session = sessionmaker(bind=engine, autoflush=False, expire_on_commit=False)()
    storage = InMemoryObjectStorage(get_settings())

    try:
        story_session = StorySession(working_title="Lantern Harbor")
        db_session.add(story_session)
        db_session.flush()

        first_story_location = storage.paths.export_asset(
            session_id=story_session.id,
            export_kind="story",
            export_id="accepted-manuscript-v1",
            extension="md",
        )
        storage.upload_text(
            first_story_location,
            "# Chapter 1\n\nMira carried the lantern home.",
            content_type="text/markdown; charset=utf-8",
        )
        first_story_asset = SessionAsset(
            session_id=story_session.id,
            asset_kind=AssetKind.STORY_TEXT,
            status=AssetStatus.SUPERSEDED,
            storage_bucket=first_story_location.bucket,
            object_path=first_story_location.key,
            mime_type="text/markdown",
        )
        db_session.add(first_story_asset)
        db_session.flush()

        second_story_location = storage.paths.export_asset(
            session_id=story_session.id,
            export_kind="story",
            export_id="accepted-manuscript-v2",
            extension="md",
        )
        storage.upload_text(
            second_story_location,
            "# Chapter 1\n\nMira carried the lantern home.\n\n# Chapter 2\n\nThe harbor slept.",
            content_type="text/markdown; charset=utf-8",
        )
        second_story_asset = SessionAsset(
            session_id=story_session.id,
            asset_kind=AssetKind.STORY_TEXT,
            status=AssetStatus.READY,
            storage_bucket=second_story_location.bucket,
            object_path=second_story_location.key,
            mime_type="text/markdown",
        )
        db_session.add(second_story_asset)
        db_session.commit()

        service = StoryDocxExportService(db_session, object_storage=storage)

        stale_export = service.ensure_docx_asset(story_session.id)
        stale_export.metadata_json = {
            **dict(stale_export.metadata_json or {}),
            "source_story_asset_id": first_story_asset.id,
        }
        db_session.commit()

        refreshed_export = service.ensure_docx_asset(story_session.id)

        assert refreshed_export.id == stale_export.id
        assert refreshed_export.metadata_json["source_story_asset_id"] == second_story_asset.id

        exported_bytes = storage.download_bytes(
            storage.paths.export_asset(
                session_id=story_session.id,
                export_kind="docx",
                export_id="final-manuscript",
                extension="docx",
            )
        )
        document = Document(BytesIO(exported_bytes))
        assert any(paragraph.text == "Chapter 2" for paragraph in document.paragraphs)
        assert any(paragraph.text == "The harbor slept." for paragraph in document.paragraphs)
    finally:
        db_session.close()
        engine.dispose()
